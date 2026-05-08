"""
Generate startup cost estimate document (DOCX) for the app studio.
Covers LLC, Plainly, Health-AI, domains, hosting, APIs, legal, and tooling.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TODAY = "March 2026"
COMPANY = "[COMPANY NAME] LLC"


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Georgia'


def add_subtitle(doc, text, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Georgia'
    if color:
        run.font.color.rgb = color


def add_section_header(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a1a2e')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.name = 'Georgia'
    return p


def add_note(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1a1a2e')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Georgia'
                    run.font.size = Pt(10)
                if c_idx == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternate row shading
            if r_idx % 2 == 1:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'f5f5f5')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shd)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph()
    return table


def add_total_row(table, label, lean, moderate, funded):
    row_cells = table.add_row().cells
    data = [label, lean, moderate, funded]
    for i, cell in enumerate(row_cells):
        cell.text = str(data[i])
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1a1a2e')
        cell._tc.get_or_add_tcPr().append(shd)


def make_cost_doc():
    doc = new_doc()

    add_title(doc, "STARTUP COST ESTIMATE")
    add_subtitle(doc, f"{COMPANY} — App Studio")
    add_subtitle(doc, f"Plainly  ·  Health-AI  ·  Future Products")
    add_subtitle(doc, f"Prepared: {TODAY}", color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Estimates based on current market pricing (March 2026). "
        "Three scenarios: Lean (bootstrap/solo founder), Moderate (small team, paid tools), "
        "Funded (early investment, full toolset). Costs shown are USD."
    )
    run.font.size = Pt(9.5)
    run.font.name = 'Georgia'
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── SECTION 1: ONE-TIME COSTS ──────────────────────────────────────────

    add_section_header(doc, "Section 1 — One-Time Startup Costs")
    add_note(doc, "These are paid once at formation and launch.")

    headers = ["Item", "Notes", "Lean", "Moderate", "Funded"]
    rows = [
        ["LLC Formation Filing Fee",
         "Wyoming ($100) or Texas ($300). Delaware $90 but adds complexity.",
         "$100", "$100", "$300"],
        ["Registered Agent (Year 1)",
         "Northwest ($125/yr) or similar. Required in state of formation.",
         "$125", "$125", "$299"],
        ["EIN (Employer ID Number)",
         "Free — obtained directly from IRS online.",
         "$0", "$0", "$0"],
        ["Operating Agreement Drafting",
         "DIY from template vs. startup attorney review.",
         "$0", "$500", "$1,500"],
        ["Google Play Developer Account",
         "One-time fee. One account covers all apps.",
         "$25", "$25", "$25"],
        ["Apple Developer Program (Year 1)",
         "$99/year. One account covers all apps.",
         "$99", "$99", "$99"],
        ["Domain — Plainly (.com or .app)",
         "plainly.com (if available) or plainly.app. ~$15–$20/yr for .com.",
         "$20", "$20", "$20"],
        ["Domain — Health-AI (.com or .app)",
         "healthai.app or similar. .app domains ~$15–$20/yr.",
         "$20", "$20", "$20"],
        ["Domain — Company / Studio site",
         "Main LLC website for investor/press landing page.",
         "$15", "$15", "$15"],
        ["Logo & Brand Design (per app)",
         "Fiverr/freelance vs. boutique studio. 2 apps.",
         "$200", "$800", "$3,000"],
        ["Business Bank Account Setup",
         "Mercury — free to open, no minimums.",
         "$0", "$0", "$0"],
        ["DBA Registration (per app, per state)",
         "Registering Plainly and Health-AI as trade names. ~$25–$100/state.",
         "$100", "$100", "$200"],
    ]

    add_table(doc, headers, rows, col_widths=[2.2, 2.5, 0.7, 0.8, 0.8])

    add_note(doc, "⬆ Lean one-time total: ~$704  |  Moderate: ~$1,804  |  Funded: ~$5,478")

    # ── SECTION 2: ANNUAL RECURRING ───────────────────────────────────────

    add_section_header(doc, "Section 2 — Annual Recurring Costs (Fixed)")
    add_note(doc, "These recur every year regardless of user volume.")

    rows2 = [
        ["Apple Developer Program",
         "$99/yr auto-renew. Covers both apps.",
         "$99", "$99", "$99"],
        ["Registered Agent",
         "Annual renewal.",
         "$125", "$125", "$299"],
        ["Domain — Plainly",
         "Annual renewal.",
         "$20", "$20", "$20"],
        ["Domain — Health-AI",
         "Annual renewal.",
         "$20", "$20", "$20"],
        ["Domain — Company site",
         "Annual renewal.",
         "$15", "$15", "$15"],
        ["General Liability Insurance",
         "App/software companies. ~$700–$1,300/yr.",
         "$700", "$900", "$1,300"],
        ["Privacy Policy Tool (Iubenda/Termly)",
         "Generates GDPR/CCPA-compliant policies. Per app or bundled.",
         "$144", "$240", "$360"],
        ["LLC Annual Report / State Fee",
         "Wyoming $60/yr; Texas $0; Delaware $300/yr.",
         "$60", "$60", "$300"],
        ["Accounting Software",
         "Wave (free) vs. QuickBooks Simple Start ($360/yr) vs. Bench ($500+/mo).",
         "$0", "$360", "$1,200"],
        ["Figma (Design)",
         "Free tier covers solo founder. Pro = $192/yr per seat.",
         "$0", "$192", "$384"],
        ["Business email (Google Workspace)",
         "Starter $72/yr per user. Covers plainly.app and healthai.app domains.",
         "$72", "$144", "$288"],
    ]

    add_table(doc, headers, rows2, col_widths=[2.2, 2.5, 0.7, 0.8, 0.8])
    add_note(doc, "⬆ Lean annual fixed: ~$1,255  |  Moderate: ~$2,175  |  Funded: ~$4,285")

    # ── SECTION 3: MONTHLY INFRASTRUCTURE ────────────────────────────────

    add_section_header(doc, "Section 3 — Monthly Infrastructure Costs (at Launch)")
    add_note(doc,
             "These are per-month costs for running both apps. "
             "'At launch' = low traffic (0–1,000 MAU). Costs scale with users.")

    headers3 = ["Service", "What It's For", "Lean/mo", "Moderate/mo", "Funded/mo"]
    rows3 = [
        ["PostgreSQL Database (Supabase/Railway)",
         "User data, logs, progress for both apps. Supabase free → $25/mo Pro.",
         "$0", "$25", "$50"],
        ["Backend Hosting — Node.js (Railway/Render)",
         "API server for both apps. Railway usage-based; Render $7–$19/mo fixed.",
         "$5", "$19", "$50"],
        ["Auth (Clerk or Supabase Auth)",
         "User auth for both apps. Free up to 50K MAU.",
         "$0", "$0", "$25"],
        ["Anthropic API — Plainly AI Coach",
         "Claude API for coach conversations. ~$50–$200/mo at low volume.",
         "$50", "$100", "$200"],
        ["Anthropic API — Health-AI Insights",
         "Claude API for health insights generation.",
         "$30", "$75", "$150"],
        ["Web Hosting — Plainly website (Vercel)",
         "Marketing/landing page. Vercel free tier covers most launches.",
         "$0", "$0", "$20"],
        ["Web Hosting — Health-AI website (Vercel)",
         "Marketing/landing page.",
         "$0", "$0", "$20"],
        ["Push Notifications (Expo / OneSignal)",
         "In-app nudges and reminders. Free tier covers early stage.",
         "$0", "$0", "$19"],
        ["Error Monitoring (Sentry)",
         "Crash reporting for both apps. Free tier: 5K errors/mo.",
         "$0", "$26", "$89"],
        ["CDN / Storage (Cloudflare R2 or S3)",
         "App assets, images. Minimal at launch.",
         "$0", "$5", "$20"],
        ["SSL Certificates",
         "Included free with Vercel, Render, Railway.",
         "$0", "$0", "$0"],
    ]

    add_table(doc, headers3, rows3, col_widths=[2.2, 2.5, 0.7, 0.8, 0.8])
    add_note(doc, "⬆ Lean monthly infra: ~$85/mo ($1,020/yr)  |  Moderate: ~$250/mo ($3,000/yr)  |  Funded: ~$643/mo ($7,716/yr)")

    # ── SECTION 4: PLAID (PLAINLY ONLY) ───────────────────────────────────

    add_section_header(doc, "Section 4 — Plaid API Costs (Plainly Only)")
    add_note(doc,
             "Plaid is required for Plainly's bank connection feature. "
             "This is the highest variable cost in the stack — it scales directly with user connections.")

    rows4 = [
        ["Plaid — Development/Sandbox",
         "Free. Use while building and testing.",
         "$0", "$0", "$0"],
        ["Plaid — Production (Launch, 0–100 active connections)",
         "~$0.50–$2.00 per connection. Estimated 50 active connections at launch.",
         "$25–$100", "$25–$100", "$100–$200"],
        ["Plaid — Production (500 connections)",
         "Volume discounts kick in around 10K+ connections.",
         "~$250–$1,000", "~$250–$1,000", "~$500–$1,500"],
        ["Plaid — Production (5,000 connections)",
         "Custom contract pricing at scale.",
         "~$2,500+", "~$2,500+", "~$2,500+"],
    ]

    headers4 = ["Plaid Tier", "Notes", "Lean", "Moderate", "Funded"]
    add_table(doc, headers4, rows4, col_widths=[2.2, 2.5, 0.7, 0.8, 0.8])
    add_note(doc,
             "Strategy: Launch with Plaid in sandbox during beta. "
             "Move to production only when onboarding paying users. "
             "Consider making bank connection optional at launch to delay Plaid costs.")

    # ── SECTION 5: DEVELOPMENT TOOLS ──────────────────────────────────────

    add_section_header(doc, "Section 5 — Development Tooling")
    add_note(doc, "Optional tools that improve velocity. Most have free tiers sufficient for early stage.")

    rows5 = [
        ["Expo EAS (Build & Submit)",
         "Cloud builds for iOS/Android. Free tier covers ~30 builds/mo.",
         "$0", "$0", "$99/mo"],
        ["GitHub",
         "Version control. Free for private repos (unlimited).",
         "$0", "$0", "$0"],
        ["Linear (Project Management)",
         "Issue tracking. Free tier covers small teams.",
         "$0", "$0", "$8/user/mo"],
        ["Notion (Docs/Wiki)",
         "Free for personal use; $10/user/mo team.",
         "$0", "$0", "$10/mo"],
        ["Postman (API Testing)",
         "Free tier is sufficient for most startups.",
         "$0", "$0", "$0"],
        ["Analytics — Mixpanel or PostHog",
         "User behavior analytics. PostHog free (self-host or cloud 1M events/mo).",
         "$0", "$0", "$0–$50/mo"],
    ]

    headers5 = ["Tool", "Notes", "Lean", "Moderate", "Funded"]
    add_table(doc, headers5, rows5, col_widths=[2.2, 2.5, 0.7, 0.8, 0.8])

    # ── SECTION 6: YEAR 1 TOTAL SUMMARY ───────────────────────────────────

    add_section_header(doc, "Section 6 — Year 1 Total Cost Summary")

    headers6 = ["Cost Category", "Lean (Bootstrap)", "Moderate", "Funded"]
    rows6 = [
        ["One-time formation & launch costs", "$704", "$1,804", "$5,478"],
        ["Annual fixed costs (recurring)", "$1,255", "$2,175", "$4,285"],
        ["Infrastructure — 12 months", "$1,020", "$3,000", "$7,716"],
        ["Plaid API — 12 months (Plainly)", "$300–$1,200", "$300–$1,200", "$1,200–$6,000"],
        ["Developer tools — 12 months", "$0", "$0", "$1,400"],
        ["Google Play (one-time, included above)", "—", "—", "—"],
    ]
    t = add_table(doc, headers6, rows6, col_widths=[2.7, 1.3, 1.1, 1.1])
    add_total_row(t, "ESTIMATED YEAR 1 TOTAL",
                  "$3,279–$4,179", "$7,279–$8,179", "$20,079–$24,879")

    doc.add_paragraph()
    add_note(doc,
             "Note: These estimates assume solo founder (Lean/Moderate) or small team of 2–3 (Funded). "
             "They do NOT include contractor development costs, marketing spend, or founder salaries. "
             "App Store and Google Play commissions (15–30%) are taken from revenue, not upfront costs.")

    # ── SECTION 7: APP STORE ECONOMICS ────────────────────────────────────

    add_section_header(doc, "Section 7 — App Store Commission Structure")
    add_note(doc,
             "Not upfront costs — commissions are deducted from revenue. "
             "Important to understand before pricing your subscription tiers.")

    rows7 = [
        ["Apple App Store — First $1M/yr revenue",
         "Small Business Program. 15% commission on all purchases and subscriptions.",
         "15%"],
        ["Apple App Store — Above $1M/yr revenue",
         "Standard rate applies once you exceed $1M in annual proceeds.",
         "30%"],
        ["Google Play Store — First $1M/yr revenue",
         "Reduced commission program matches Apple's Small Business rates.",
         "15%"],
        ["Google Play Store — Above $1M/yr revenue",
         "Standard rate applies.",
         "30%"],
        ["Auto-renewing subscriptions (both stores)",
         "15% rate applies from day 1 regardless of revenue tier.",
         "15%"],
    ]

    headers7 = ["Tier", "Notes", "Commission"]
    add_table(doc, headers7, rows7, col_widths=[2.2, 3.5, 1.0])

    add_note(doc,
             "Pricing implication: If you charge $9.99/month for a subscription, "
             "you net ~$8.49 after the 15% commission. Plan your unit economics accordingly.")

    # ── SECTION 8: COST REDUCTION STRATEGIES ─────────────────────────────

    add_section_header(doc, "Section 8 — Cost Reduction Strategies")

    strategies = [
        ("Launch Plainly's bank connection as optional",
         "Delays Plaid production costs until you have paying users. Use manual entry first."),
        ("Use Supabase for both DB and Auth",
         "Eliminates need for separate auth service (Clerk). One platform handles both."),
        ("Deploy on Railway at launch",
         "Usage-based pricing means near-zero cost until you have real traffic."),
        ("Expo EAS free tier for builds",
         "Free tier covers ~30 builds/month — sufficient for early launch cycles."),
        ("Use PostHog for analytics (self-host or free cloud)",
         "Replaces paid analytics tools. 1M events/month free on cloud."),
        ("Wyoming LLC over Delaware",
         "Wyoming: $100 formation + $60/yr. Delaware: $90 formation + $300/yr franchise tax."),
        ("Batch Anthropic API calls",
         "Batch API offers 50% discount on tokens. Use for non-real-time insight generation."),
        ("Single Apple Developer account",
         "$99/year covers unlimited apps. No need for separate accounts per product."),
        ("Vercel free tier for marketing sites",
         "Both app marketing pages can run free on Vercel's hobby tier."),
        ("Defer general liability insurance",
         "Get it before launch, but shop around — app companies often qualify for lower rates."),
    ]

    for title, detail in strategies:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(10.5)
        run2 = p.add_run(detail)
        run2.font.name = 'Georgia'
        run2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    # ── FOOTER NOTE ───────────────────────────────────────────────────────

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Pricing sources: Apple Developer, Google Play Console, Supabase, Railway, Render, "
        "Vercel, Clerk, Plaid, Anthropic, Sentry, OneSignal, Mercury, Figma, Iubenda, "
        "Wyoming/Texas/Delaware SOS — all current as of March 2026. Prices subject to change."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save("/home/user/Claude-/docs/startup_cost_estimate.docx")
    print("✓ Startup Cost Estimate generated")


if __name__ == "__main__":
    make_cost_doc()
