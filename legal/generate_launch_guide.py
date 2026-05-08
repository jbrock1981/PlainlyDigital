"""
Generate phased launch guide DOCX — LLC first, then Health-AI, then Plainly,
then future apps. Step-by-step with cost at each step and running totals.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TODAY = "March 2026"
DARK = "1a1a2e"
TEAL = "00b4d8"
GREEN = "2d6a4f"
ORANGE = "e76f51"
PURPLE = "6d3a9c"


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)
    doc.styles['Normal'].font.name = 'Georgia'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Georgia'
    run.font.color.rgb = hex_to_rgb(DARK)


def add_centered(doc, text, size=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Georgia'
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = hex_to_rgb(color)
    return p


def add_phase_header(doc, phase_num, title, subtitle, color_hex):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

    run1 = p.add_run(f"PHASE {phase_num}  ")
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.name = 'Georgia'
    run1.font.color.rgb = hex_to_rgb(color_hex)

    run2 = p.add_run(f"—  {title.upper()}")
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.name = 'Georgia'
    run2.font.color.rgb = hex_to_rgb(DARK)

    p2 = doc.add_paragraph(subtitle)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    for run in p2.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = hex_to_rgb("666666")

    # Colored rule
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(8)
    pPr = p3._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_step(doc, step_num, title, description, cost, notes=None, optional=False):
    """Add one numbered step with cost."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0)

    # Step number badge
    run_num = p.add_run(f"  {step_num:02d}  ")
    run_num.bold = True
    run_num.font.name = 'Courier New'
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = hex_to_rgb("ffffff")

    run_title = p.add_run(f"  {title}")
    run_title.bold = True
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(11)
    if optional:
        run_opt = p.add_run("  [OPTIONAL]")
        run_opt.font.name = 'Georgia'
        run_opt.font.size = Pt(9)
        run_opt.font.color.rgb = hex_to_rgb("999999")

    # Description
    p2 = doc.add_paragraph(description)
    p2.paragraph_format.left_indent = Inches(0.35)
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(2)
    for run in p2.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(10.5)

    # Cost line
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.35)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(2)
    run_label = p3.add_run("Cost: ")
    run_label.bold = True
    run_label.font.name = 'Georgia'
    run_label.font.size = Pt(10)
    run_cost = p3.add_run(cost)
    run_cost.font.name = 'Georgia'
    run_cost.font.size = Pt(10)
    run_cost.font.color.rgb = hex_to_rgb(GREEN)

    # Notes
    if notes:
        p4 = doc.add_paragraph(f"↳  {notes}")
        p4.paragraph_format.left_indent = Inches(0.35)
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(6)
        for run in p4.runs:
            run.font.name = 'Georgia'
            run.font.size = Pt(9.5)
            run.italic = True
            run.font.color.rgb = hex_to_rgb("555555")
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_phase_total(doc, phase_label, one_time, monthly, annual_note, running_total, color_hex):
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    labels = [phase_label, f"One-time: {one_time}", f"Monthly: {monthly}", f"Running total (Yr 1): {running_total}"]
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = labels[i]
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shd)

    if annual_note:
        p = doc.add_paragraph(f"  {annual_note}")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = 'Georgia'
            run.font.size = Pt(9.5)
            run.italic = True
            run.font.color.rgb = hex_to_rgb("555555")

    doc.add_paragraph()


def make_guide():
    doc = new_doc()

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_title(doc, "LAUNCH ROADMAP & COST GUIDE")
    add_centered(doc, "Step-by-Step — From LLC to Multi-App Studio", size=12, italic=True, color="666666")
    add_centered(doc, "Health-AI  →  Plainly  →  Future Products", size=11, bold=True, color=DARK)
    add_centered(doc, f"Prepared {TODAY}", size=9, italic=True, color="999999")
    doc.add_paragraph()

    add_centered(doc,
        "Each step is numbered, sequenced, and priced. Complete each step before moving to the next. "
        "Costs shown are the lean/bootstrap scenario. Running totals reflect cumulative Year 1 spend.",
        size=10, italic=True, color="555555")
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 0 — LLC & FOUNDATION
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, 0, "LLC & Business Foundation",
        "Do this once. Everything else builds on it.", DARK)

    step = 1
    add_step(doc, step, "Choose Your State & Form the LLC",
        "File your Articles of Organization with the Secretary of State. "
        "Wyoming is the best lean choice: $100 to file, $60/yr to maintain, no state income tax, "
        "strong privacy protections. Texas is also good ($300 file, $0/yr ongoing). Avoid California ($800/yr).",
        "Wyoming: $100 one-time  |  Texas: $300 one-time",
        "File online at the Wyoming SOS website (wyoming.gov). Takes 1–3 business days.")

    step += 1
    add_step(doc, step, "Appoint a Registered Agent",
        "Every LLC needs a registered agent with a physical address in your state of formation. "
        "They receive legal documents on your behalf. Northwest Registered Agent is the most trusted "
        "budget option. Required by law — cannot skip.",
        "$125/year (Northwest Registered Agent) or $0 if you use your own address",
        "Your first year is often included or discounted when bundled with LLC formation tools.")

    step += 1
    add_step(doc, step, "Get Your EIN (Employer Identification Number)",
        "Your LLC's federal tax ID — like a Social Security number for the business. "
        "Used to open a bank account, file taxes, and set up payroll later. "
        "Apply directly on the IRS website (irs.gov → Apply for EIN). Free, instant.",
        "$0 — completely free from IRS.gov",
        "Takes 5 minutes online. You'll get it immediately. Save the confirmation PDF.")

    step += 1
    add_step(doc, step, "Draft or Purchase an Operating Agreement",
        "This is the internal contract that defines ownership, profit splits, voting rights, "
        "and what happens if a founder leaves. Even as a sole member, most banks require it. "
        "Single-member LLCs can use a template; multi-member LLCs should have an attorney review.",
        "Solo/template: $0  |  Attorney-reviewed: $500–$1,500",
        "Sites like Rocket Lawyer or Clerky have solid templates. If you have co-founders, spend the money on an attorney.")

    step += 1
    add_step(doc, step, "Open a Business Bank Account",
        "Keep business and personal finances 100% separate from day one. This protects your "
        "LLC liability shield and simplifies taxes enormously. Mercury is the best option for "
        "startups: no monthly fees, no minimums, free wire transfers, and built for tech companies.",
        "$0 — Mercury is free (mercury.com)",
        "You'll need your EIN and Articles of Organization to open the account. Takes 1–3 business days for approval.")

    step += 1
    add_step(doc, step, "Set Up Business Email",
        "Create professional email addresses for the business using Google Workspace. "
        "You'll need a domain for this — grab your company/studio domain now even if the "
        "app-specific domains come later. Example: hello@yourstudio.com",
        "$72/year ($6/mo) for 1 user on Google Workspace Starter",
        "Register your studio domain at Namecheap or Cloudflare (~$12–15/year). Do this before the bank account if possible.")

    step += 1
    add_step(doc, step, "Register Your Studio/Company Domain",
        "This is the domain for your LLC itself — not the app domains (those come later). "
        "Example: mystudio.com or yourname.app. Used for your main website, email, and investor/press pages.",
        "$12–20/year at Namecheap, Cloudflare, or Google Domains",
        "Cloudflare Registrar is the cheapest option with no markup over ICANN cost.")

    step += 1
    add_step(doc, step, "Set Up Accounting (Bookkeeping)",
        "Even before you have revenue, track every business expense from day one. "
        "This makes tax time manageable and lets you deduct startup costs. "
        "Wave Accounting is free and solid for solo founders. QuickBooks if you want more automation.",
        "Wave: $0  |  QuickBooks Simple Start: $30/month ($360/yr)",
        "Log every receipt from Step 1 onward. Your formation fees, registered agent, and domains are all deductible startup costs.")

    add_phase_total(doc,
        "PHASE 0 COMPLETE",
        "$307–$537 (one-time)",
        "$16–$52/mo (email + accounting)",
        "Running year 1 fixed: $197–$624. You now have a legal entity, bank account, and email.",
        "~$504–$1,161",
        DARK)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 1 — HEALTH-AI APP
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, 1, "Health-AI — Build & Launch",
        "Your first app. Establish all shared infrastructure here — the next app reuses most of it.", TEAL)

    add_step(doc, step := step + 1, "Register the Health-AI Domain",
        "Secure your app's domain name. Options: healthai.app, gethealthai.com, tryhealthai.com. "
        "The .app TLD is HTTPS-enforced by browsers and signals a modern product. "
        "Check availability and register before you announce anything.",
        "$15–$20/year for .app or .com domain",
        "Use Cloudflare Registrar for cheapest renewal rates. Also grab the .com variant if the .app is your primary.")

    add_step(doc, step := step + 1, "Set Up Developer Accounts (Apple + Google)",
        "One Apple Developer account and one Google Play account covers ALL your apps forever. "
        "Apple: $99/year subscription. Google: $25 one-time fee. Both accounts will be used "
        "for Health-AI now and every future app you launch — you never pay this again.",
        "Apple: $99/year  |  Google: $25 one-time",
        "Use your business email when registering. For Apple, enroll as an 'Organization' not 'Individual' if you have an LLC with a DUNS number (free from Dun & Bradstreet).")

    add_step(doc, step := step + 1, "Set Up Version Control & Project Structure",
        "Create a GitHub organization for your studio (free). Set up two private repos: "
        "one for the Health-AI mobile app (Expo/React Native), one for the backend API. "
        "This is where all your code lives. Never build without version control.",
        "$0 — GitHub free tier includes unlimited private repos",
        "GitHub organization: free. If you want protected branches and PR reviews, GitHub Team is $4/user/mo.")

    add_step(doc, step := step + 1, "Initialize the Expo (React Native) Project",
        "Bootstrap the Health-AI mobile app using Expo. Expo handles cross-platform iOS/Android "
        "development with one codebase, plus over-the-air updates (EAS Update) without App Store re-approval. "
        "Free tier is sufficient to build and launch.",
        "$0 — Expo free tier",
        "Run: npx create-expo-app@latest HealthAI. Set up EAS CLI for building. Free tier includes enough build minutes for launch.")

    add_step(doc, step := step + 1, "Set Up Authentication (Supabase Auth)",
        "Every app needs user accounts. Supabase provides auth + database in one platform. "
        "Free tier covers 50,000 monthly active users — more than enough for launch. "
        "Set up email/password auth first, then add Google and Apple Sign-In (required by Apple for apps that use social login).",
        "$0 free tier → $25/month Pro when you exceed 50K MAU",
        "Supabase free tier: 50K MAU, 500MB storage, 2 projects. Perfect for launch.")

    add_step(doc, step := step + 1, "Set Up the PostgreSQL Database (Supabase)",
        "Your Supabase project includes a PostgreSQL database automatically. "
        "Design your schema for Health-AI: users, health_logs, goals, ai_conversations, insights. "
        "No separate database hosting needed — it's bundled with your Supabase auth.",
        "$0 on Supabase free tier (included with auth setup above)",
        "You get 500MB of storage free. For a new app this is plenty. Upgrade to Pro ($25/mo) when you need more.")

    add_step(doc, step := step + 1, "Set Up the Backend API (Node.js on Railway)",
        "Deploy a Node.js/Express (or Fastify) backend on Railway. This handles your business logic, "
        "Claude API calls, and acts as the secure intermediary between the app and Supabase. "
        "Railway charges by usage — expect near-zero cost until you have real traffic.",
        "$0–$5/month on Railway (usage-based, near-zero at low traffic)",
        "Alternative: Render.com at $7/month fixed. Railway is cheaper at low scale; Render is more predictable.")

    add_step(doc, step := step + 1, "Set Up Anthropic API (Claude) for AI Features",
        "Create an Anthropic account, get an API key, and integrate it into your backend. "
        "Health-AI will use Claude for health insight generation, goal coaching, and trend analysis. "
        "Use Claude Haiku for lightweight responses ($1/M input tokens) and Sonnet for complex analysis.",
        "$30–$80/month at launch (estimated, based on usage)",
        "Never call the Anthropic API directly from the mobile app — always route through your backend so your API key stays secret.")

    add_step(doc, step := step + 1, "Build the App (Development Phase)",
        "This is where the bulk of your time (and potential contractor costs) goes. "
        "Core screens: onboarding, dashboard, health log entry, AI insights, goals, settings. "
        "If building solo, budget 2–4 months of focused development. "
        "If hiring a contractor, expect $50–$150/hour for a React Native developer.",
        "Solo: $0 labor  |  Freelancer: $5,000–$20,000  |  Agency: $30,000–$80,000",
        "Fiverr/Toptal for vetted freelancers. Build an MVP first — ship the core loop before adding polish.")

    add_step(doc, step := step + 1, "Brand Design — Health-AI",
        "App icon, splash screen, color palette, typography, and basic UI component library. "
        "The app icon is critical — it's your first impression in the App Store. "
        "Figma is free for solo founders. For the icon and brand assets, consider a freelancer.",
        "DIY with Figma: $0  |  Freelance designer: $200–$800  |  Studio: $2,000–$5,000",
        "Figma free tier works perfectly for solo founders. Upgrade to Pro ($16/mo) only if you have a team.")

    add_step(doc, step := step + 1, "Set Up Push Notifications",
        "Health apps live and die by engagement. Set up Expo push notifications for daily "
        "health check-in reminders, AI insight delivery, and goal milestone alerts. "
        "Expo's built-in push service is free and works on both iOS and Android.",
        "$0 — Expo push notifications are free",
        "If you need advanced segmentation or A/B testing of notifications, OneSignal free tier handles it.")

    add_step(doc, step := step + 1, "Set Up Error Monitoring (Sentry)",
        "Before you launch, you need to know when things break. Sentry captures crashes, "
        "errors, and performance issues from both the mobile app and backend. "
        "Free tier covers 5,000 errors/month — more than enough for launch.",
        "$0 on Sentry free tier",
        "Integrate Sentry into both your Expo app and Node.js backend. Takes about 30 minutes to set up.")

    add_step(doc, step := step + 1, "Generate Privacy Policy & Terms of Service",
        "App Store (both Apple and Google) require a privacy policy link. "
        "GDPR and CCPA compliance is also mandatory if you have users in the EU or California. "
        "Health data is sensitive — your privacy policy must specifically address health data handling.",
        "$72–$120/year (Iubenda or Termly)",
        "Iubenda Essentials: $72/yr. Termly Starter: $120/yr. Both generate App Store-compliant policies. "
        "Health data falls under HIPAA consideration — consult an attorney if you store identifiable health records.")

    add_step(doc, step := step + 1, "Set Up the App Marketing Website",
        "Before App Store submission, you need a landing page with: app description, screenshots, "
        "privacy policy link, and support email. This is your App Store 'support URL'. "
        "Deploy on Vercel free tier using your healthai.app domain.",
        "$0 — Vercel free tier + domain already purchased in Step 9",
        "A simple one-page site is fine for launch. Use Framer, Webflow, or just Next.js deployed on Vercel.")

    add_step(doc, step := step + 1, "TestFlight Beta & Internal Testing",
        "Before submitting to the App Store, distribute through TestFlight (iOS) and Google Play "
        "Internal Testing (Android) to get real feedback. TestFlight is free, included in your "
        "Apple Developer account. Get 10–20 beta users minimum before public launch.",
        "$0 — included in your $99 Apple Developer account",
        "Build your TestFlight IPA using EAS Build: eas build --platform ios --profile preview. First build takes ~10 minutes.")

    add_step(doc, step := step + 1, "Submit to App Store & Google Play",
        "Submit Health-AI for App Store Review (Apple) and Google Play Review. "
        "Apple review typically takes 1–3 days. Google Play is usually 1–7 days for first submission. "
        "Prepare: app screenshots (all required sizes), description, keywords, and support URL.",
        "$0 — covered by your existing developer accounts",
        "Apple requires screenshots for iPhone 6.9\", 6.5\", and iPad Pro 13\". Use EAS Submit to automate the upload process.")

    add_step(doc, step := step + 1, "Set Up Analytics",
        "Track what users actually do in your app. PostHog is the best free option — "
        "1 million events/month free on their cloud tier. Track: onboarding completion rate, "
        "daily active users, AI insight generation rate, retention D1/D7/D30.",
        "$0 — PostHog free tier (1M events/month)",
        "PostHog is open-source and self-hostable if you later need more control. Avoid Firebase Analytics — harder to query.")

    add_phase_total(doc,
        "PHASE 1 COMPLETE — HEALTH-AI LAUNCHED",
        "$349–$21,099 (varies on dev costs)",
        "$45–$115/mo infrastructure",
        "Year 1 infra: ~$540–$1,380. Running total includes Phase 0. Dev labor not included in infrastructure estimate.",
        "~$1,353–$3,165 (infra only, no dev labor)",
        TEAL)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 2 — PLAINLY (FINANCIAL LITERACY)
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, 2, "Plainly — Financial Literacy App",
        "Your second app. Most infrastructure from Phase 1 is already in place — "
        "costs are significantly lower for app #2.", ORANGE)

    step_start = step

    add_step(doc, step := step + 1, "Register the Plainly Domain",
        "Secure plainly.app, plainly.com, getplainly.com, or similar. "
        "Check availability at Namecheap or Cloudflare. Register both .com and .app "
        "if the name is important to you long-term.",
        "$15–$40/year (one or two domains)",
        "If plainly.com is taken, consider tryplainly.com or getplainly.com. Keep the name short and memorable.")

    add_step(doc, step := step + 1, "Create Plainly Supabase Project",
        "Create a second Supabase project for Plainly's data — separate from Health-AI "
        "for data isolation and cleaner architecture. Free tier covers 2 projects, so "
        "this is still free. Schema: users, financial_goals, lessons, quiz_results, ai_conversations.",
        "$0 — second project on Supabase free tier",
        "Supabase free tier allows 2 active projects. If you're already using 2, you'll need to upgrade one account to Pro ($25/mo).")

    add_step(doc, step := step + 1, "Initialize Plainly Expo Project & Repo",
        "Bootstrap the Plainly app using the same Expo setup from Phase 1. "
        "Create a new GitHub repo and Expo project. You already know the stack — "
        "this setup step takes a fraction of the time it did for Health-AI.",
        "$0 — reuses existing Expo and GitHub setup",
        "Consider creating a shared component library (monorepo) to reuse UI components across both apps.")

    add_step(doc, step := step + 1, "Set Up Plaid API (Bank Connection — Sandbox First)",
        "Plainly's core feature is connecting to users' bank accounts for real financial tracking. "
        "Plaid is the industry standard. Start in Sandbox mode (free) while building. "
        "Only move to Production when you're ready for real users.",
        "Sandbox: $0  |  Production: ~$0.50–$2.00 per active connection/month",
        "Apply for Plaid Production access early — it requires a review process that can take 1–2 weeks. "
        "Consider launching with manual entry first to avoid blocking on Plaid approval.")

    add_step(doc, step := step + 1, "Build Plainly Core Features",
        "Core screens: financial literacy lessons, quiz/assessment, goal setting, budget overview, "
        "AI financial coach (Claude), bank connection (Plaid), progress tracking, and badge/achievement system. "
        "Because you already built the auth, backend, and AI integration for Health-AI, "
        "you can reuse those patterns — this app goes faster.",
        "Solo: $0 labor  |  Freelancer: $4,000–$15,000",
        "The AI coach and lesson content are the most time-intensive parts. Consider using Claude to help generate the financial literacy curriculum content itself.")

    add_step(doc, step := step + 1, "Brand Design — Plainly",
        "Separate visual identity from Health-AI. Plainly should feel trustworthy, clear, "
        "and approachable — think clean blues and greens, not medical. "
        "App icon, splash screen, color system, and UI components.",
        "DIY: $0  |  Freelancer: $200–$800",
        "If you used the same designer from Phase 1, they'll understand your aesthetic and can work faster.")

    add_step(doc, step := step + 1, "Generate Privacy Policy & Terms — Plainly",
        "Plainly handles financial data and bank connections — your privacy policy must "
        "explicitly cover this. Iubenda or Termly let you create app-specific policies. "
        "Financial data and Plaid usage requires additional disclosures.",
        "$72–$120/year additional (separate policy for Plainly)",
        "If you're already on a paid Iubenda/Termly plan, you may be able to add additional sites/apps at a reduced rate.")

    add_step(doc, step := step + 1, "Set Up Plainly Marketing Website",
        "Same approach as Health-AI: one-page landing site deployed on Vercel free tier "
        "at your Plainly domain. App screenshots, description, privacy policy link, "
        "and an email waitlist if you want to build pre-launch interest.",
        "$0 — Vercel free tier + domain purchased in Step above",
        "Consider adding an email waitlist (Mailchimp free or Loops.so) before launch to capture early interest.")

    add_step(doc, step := step + 1, "Move Plaid to Production & Test Bank Connections",
        "Once your Plaid production application is approved and the app is stable, "
        "switch from Sandbox to Production. Test real bank connections with your own account first. "
        "Plaid charges per active linked item (bank connection) per month.",
        "Production start: ~$50–$200/month at 50–100 connections",
        "Plaid's production costs scale fast. Set a budget alert and consider limiting the number of bank connections per free user.")

    add_step(doc, step := step + 1, "TestFlight Beta, Submit to App Store & Google Play",
        "Same process as Health-AI. Beta test with 10–20 users on TestFlight and Google Play "
        "Internal Testing. Collect feedback, fix bugs, then submit for public review.",
        "$0 — already covered by your existing developer accounts",
        "You now have two live apps on both stores under the same developer accounts. No additional fees.")

    add_phase_total(doc,
        "PHASE 2 COMPLETE — PLAINLY LAUNCHED",
        "$302–$16,000 (varies on dev labor)",
        "+ $50–$250/mo for Plaid (scales with users)",
        "Year 1 Plaid estimate at low usage: +$600–$3,000. Two apps live in both stores.",
        "~$2,255–$6,765 (infra only, cumulative)",
        ORANGE)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 3 — FUTURE APPS
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, 3, "Future Apps — Incremental Cost Model",
        "By Phase 3, your infrastructure is fully built. Each new app costs a fraction of the first two.", PURPLE)

    add_step(doc, step := step + 1, "Register New App Domain",
        "The primary cost for each new app launch. Same process as before.",
        "$15–$40/year per app",
        "Apple and Google developer accounts already paid. Backend, auth, and DB infrastructure already running.")

    add_step(doc, step := step + 1, "Create New Supabase Project (or extend existing)",
        "Supabase free tier covers 2 projects. Once at capacity, either upgrade Pro ($25/mo) "
        "or add a new Supabase organization. You may consolidate into one Pro account by this point.",
        "$0 (if within free tier limit) or $25/month Pro",
        "Alternatively, use schema isolation within a single Supabase database to run multiple apps on one project.")

    add_step(doc, step := step + 1, "Scaffold New App from Shared Codebase",
        "By app #3, you'll have a reusable component library, auth patterns, API boilerplate, "
        "and backend patterns from the first two apps. New apps ship significantly faster.",
        "$0 — existing tooling and setup",
        "Consider building an internal Expo template or monorepo that new apps can be generated from in minutes.")

    add_step(doc, step := step + 1, "Generate App-Specific Privacy Policy",
        "Each app needs its own policy. If on a paid Iubenda or Termly plan, "
        "additional sites/apps are often included or available at a reduced rate.",
        "$0–$72/year per additional app depending on your plan",
        "Review your existing plan — multi-site plans exist that cover unlimited apps.")

    add_step(doc, step := step + 1, "Build, Test, Submit",
        "Same pipeline as Phase 1 and 2 — just faster because the patterns are established.",
        "$0 for App Store submission (covered by existing accounts)",
        "By your 3rd app, TestFlight and Play Console submission should be a well-practiced routine.")

    add_phase_total(doc,
        "PER ADDITIONAL APP (Phase 3+)",
        "$15–$112 per app (infra only)",
        "+ Marginal increase in Claude API usage",
        "Each new app adds only domain + policy costs if it uses shared backend + auth infrastructure.",
        "~$15–$112 per new app launch (beyond existing infra)",
        PURPLE)

    # ════════════════════════════════════════════════════════════════════════
    # CUMULATIVE SUMMARY TABLE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CUMULATIVE COST SUMMARY — YEAR 1 (INFRASTRUCTURE ONLY)")
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.font.color.rgb = hex_to_rgb(DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), DARK)
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()
    headers_s = ["Phase", "What You Complete", "One-Time Costs", "Monthly Add", "Yr 1 Phase Cost"]
    rows_s = [
        ["Phase 0 — LLC", "Legal entity, bank account, email, accounting", "$307–$537", "$16–$52/mo", "$499–$1,161"],
        ["Phase 1 — Health-AI", "App live on iOS + Android + web marketing site", "$349–$749", "+$45–$115/mo", "$540–$1,380"],
        ["Phase 2 — Plainly", "App live on iOS + Android, Plaid connected", "$302–$502", "+$50–$250/mo", "$902–$3,502"],
        ["Phase 3+ (per app)", "Each additional app launch", "$15–$112", "+$10–$50/mo", "$135–$712"],
    ]
    t = doc.add_table(rows=1 + len(rows_s), cols=5)
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers_s):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), DARK)
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

    for r_idx, row_data in enumerate(rows_s):
        row_cells = t.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Georgia'
                    run.font.size = Pt(10)
                if c_idx == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'f5f5f5')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shd)

    # Total row
    total_row = t.add_row()
    totals = ["YEAR 1 TOTAL (P0 + P1 + P2)", "LLC + 2 apps live", "$958–$1,788", "$111–$417/mo", "$1,941–$6,043"]
    for i, cell in enumerate(total_row.cells):
        cell.text = totals[i]
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), GREEN.lstrip('#') if not GREEN.startswith('2') else GREEN)
        cell._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph()

    # ── IMPORTANT EXCLUSIONS NOTE ──────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("What's NOT included in these costs:")
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(11)

    exclusions = [
        ("Developer/contractor labor", "The single biggest cost variable. Estimated separately above."),
        ("Marketing & user acquisition", "Paid ads, influencer partnerships, PR. This scales with your goals."),
        ("App Store commissions", "15–30% of revenue, deducted at payout. Not an upfront cost."),
        ("Taxes", "Consult a CPA. LLCs are pass-through entities — profit is taxed on your personal return."),
        ("HIPAA compliance (if applicable)",
         "If Health-AI stores identifiable health records, formal HIPAA compliance adds $2,000–$15,000+/yr."),
        ("Attorneys beyond operating agreement", "Trademark registration (~$250–$350/class), contract review, etc."),
    ]

    for title, detail in exclusions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run1 = p.add_run(f"• {title}: ")
        run1.bold = True
        run1.font.name = 'Georgia'
        run1.font.size = Pt(10.5)
        run2 = p.add_run(detail)
        run2.font.name = 'Georgia'
        run2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Pricing current as of March 2026. All costs USD. Lean scenario assumes solo founder "
        "building with free tiers where available. Source data from Apple, Google, Supabase, "
        "Railway, Anthropic, Plaid, Sentry, Iubenda, Mercury, and state SOS websites."
    )
    for run in p.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = hex_to_rgb("999999")

    doc.save("/home/user/Claude-/docs/phased_launch_guide.docx")
    print("✓ Phased launch guide generated")


if __name__ == "__main__":
    make_guide()
