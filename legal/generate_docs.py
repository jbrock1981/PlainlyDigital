"""
Generate all legal documents in DOCX format for attorney review.
Documents cover the Gen Z app studio LLC and its products: Plainly & Vinla.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = "March 14, 2026"
COMPANY = "Plainly Digital LLC"
STATE = "Tennessee"
MEMBER = "Jonathan Brock"
MEMBER_ADDRESS = "1309 Case Rd, Prospect, TN 38477"
EMAIL_PLAINLY = "legal@plainly.app"
EMAIL_VINLA = "legal@<your vinla domain>"


# ─── Helpers ────────────────────────────────────────────────────────────────

def new_doc(title):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Georgia'
    doc.add_paragraph()


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_section(doc, number, heading):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {heading}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Georgia'


def add_subsection(doc, heading):
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.bold = True
    run.font.name = 'Georgia'


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Georgia'
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def add_sig_block(doc, parties):
    doc.add_paragraph()
    add_body(doc, "IN WITNESS WHEREOF, the undersigned have executed this Agreement as of the date first written above.")
    doc.add_paragraph()
    for party in parties:
        p = doc.add_paragraph()
        run = p.add_run(party)
        run.bold = True
        p = doc.add_paragraph("Signature: ___________________________________")
        p = doc.add_paragraph("Print Name: ___________________________________")
        p = doc.add_paragraph("Title: ___________________________________")
        p = doc.add_paragraph("Date: ___________________________________")
        doc.add_paragraph()


def add_attorney_note(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(
        "ATTORNEY REVIEW NOTE: This document is a draft prepared for attorney review. "
        "Bracketed items [LIKE THIS] must be completed with accurate information before execution. "
        "This document does not constitute legal advice. Have a licensed attorney in the applicable "
        "jurisdiction review and finalize this document before signing or publishing."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)


# ─── 1. LLC Operating Agreement ─────────────────────────────────────────────

def make_operating_agreement():
    doc = new_doc("LLC Operating Agreement")
    add_title(doc, "OPERATING AGREEMENT")
    add_subtitle(doc, f"of {COMPANY}")
    add_subtitle(doc, f"A {STATE} Limited Liability Company")
    doc.add_paragraph()
    add_subtitle(doc, f"Effective Date: {TODAY}")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_section(doc, 1, "FORMATION")
    add_body(doc, f"1.1  Name. The name of the limited liability company is {COMPANY} (the \"Company\").")
    add_body(doc, f"1.2  Principal Office. The principal office of the Company shall be located at {MEMBER_ADDRESS}, or such other place as the Member may designate from time to time.")
    add_body(doc, f"1.3  Registered Agent. The Company shall maintain a registered agent in {STATE} as required by applicable law. The initial registered agent is [REGISTERED AGENT NAME AND ADDRESS].")
    add_body(doc, f"1.4  Purpose. The Company is organized to engage in any lawful act or activity for which a limited liability company may be organized under the laws of {STATE}, including but not limited to: developing, publishing, marketing, and operating mobile and web applications; generating revenue through subscriptions, advertising, and in-app purchases; entering into contracts with third-party service providers; and all activities incidental or related thereto.")
    add_body(doc, f"1.5  Term. The Company shall continue perpetually unless dissolved in accordance with this Agreement or applicable law.")

    add_section(doc, 2, "MEMBER")
    add_body(doc, f"2.1  Sole Member. The Company is a single-member LLC. The sole member is {MEMBER} (the \"Member\"), holding 100% of the membership interests.")
    add_body(doc, "2.2  Additional Members. No additional members may be admitted without an amendment to this Agreement executed by the Member.")
    add_body(doc, "2.3  Limited Liability. The Member shall not be personally liable for the debts, obligations, or liabilities of the Company solely by reason of being a member.")

    add_section(doc, 3, "CAPITAL CONTRIBUTIONS")
    add_body(doc, "3.1  Initial Contribution. The Member's initial capital contribution is set forth in Exhibit A attached hereto.")
    add_body(doc, "3.2  Additional Contributions. The Member may, but is not required to, make additional capital contributions at any time.")
    add_body(doc, "3.3  No Interest. No interest shall be paid on capital contributions.")

    add_section(doc, 4, "MANAGEMENT")
    add_body(doc, "4.1  Member-Managed. The Company shall be member-managed. The Member shall have full authority to manage the business and affairs of the Company.")
    add_body(doc, "4.2  Authority. The Member is authorized to: (a) execute contracts and agreements; (b) open and manage bank accounts; (c) hire and terminate employees and contractors; (d) incur indebtedness; (e) acquire and dispose of property; (f) make all business decisions on behalf of the Company.")
    add_body(doc, "4.3  Products and Trade Names. The Company currently operates or intends to operate the following products under the Company's ownership, using DBA (doing business as) trade names registered in accordance with applicable state law:")
    add_body(doc, "         (a) Plainly — a financial literacy and AI coaching application targeting young adults;")
    add_body(doc, "         (b) Vinla — a personal health intelligence application for logging food, water, exercise, and sleep with AI-generated insights;")
    add_body(doc, "         (c) Any additional products developed or acquired by the Company from time to time.")
    add_body(doc, "All intellectual property, revenue, and assets associated with each product are owned exclusively by the Company.")

    add_section(doc, 5, "ALLOCATIONS AND DISTRIBUTIONS")
    add_body(doc, "5.1  Allocations. All profits and losses of the Company shall be allocated 100% to the Member.")
    add_body(doc, "5.2  Distributions. Distributions shall be made to the Member at such times and in such amounts as the Member determines in their sole discretion, subject to applicable law and maintenance of adequate operating reserves.")
    add_body(doc, "5.3  Operating Reserve. The Member shall maintain a reasonable operating reserve, as determined by the Member from time to time.")

    add_section(doc, 6, "INTELLECTUAL PROPERTY")
    add_body(doc, "6.1  Company Ownership. All intellectual property created, developed, or acquired in connection with the Company's business — including source code, designs, trademarks, domain names, user data structures, AI model configurations, and marketing materials — is and shall remain the exclusive property of the Company.")
    add_body(doc, "6.2  Contractor IP Assignment. Any person or entity engaged as an independent contractor shall execute a written Intellectual Property Assignment Agreement prior to commencing work, assigning all work product to the Company.")
    add_body(doc, "6.3  Third-Party IP. The Company shall obtain appropriate licenses for all third-party intellectual property used in its products, including open-source software, APIs, and AI models.")

    add_section(doc, 7, "BANKING AND FINANCES")
    add_body(doc, "7.1  Bank Accounts. The Company shall maintain one or more business bank accounts in the Company's name. Personal and business funds shall not be commingled.")
    add_body(doc, "7.2  Fiscal Year. The Company's fiscal year shall end on December 31 of each year.")
    add_body(doc, "7.3  Books and Records. The Company shall maintain complete and accurate books of account and other records, kept at the Company's principal office.")
    add_body(doc, "7.4  Tax Treatment. The Company shall be treated as a disregarded entity for federal income tax purposes unless the Member elects otherwise. The Member may elect S-Corporation tax treatment upon meeting applicable IRS requirements.")

    add_section(doc, 8, "TRANSFER OF MEMBERSHIP INTEREST")
    add_body(doc, "8.1  Restriction. The Member may not transfer, assign, pledge, or encumber all or any portion of their membership interest without compliance with applicable law.")
    add_body(doc, "8.2  Transfer to Trust or Entity. The Member may transfer their membership interest to a revocable living trust for estate planning purposes without restriction.")

    add_section(doc, 9, "DISSOLUTION")
    add_body(doc, "9.1  Events of Dissolution. The Company shall be dissolved upon: (a) the written determination of the Member to dissolve; (b) entry of a judicial decree of dissolution; or (c) any other event requiring dissolution under applicable law.")
    add_body(doc, "9.2  Winding Up. Upon dissolution, the Member or a designated liquidating trustee shall wind up the Company's affairs, pay or provide for all debts and liabilities, and distribute remaining assets to the Member.")

    add_section(doc, 10, "INDEMNIFICATION")
    add_body(doc, "10.1  Indemnification. The Company shall indemnify and hold harmless the Member from and against any claims, liabilities, damages, costs, and expenses (including reasonable attorneys' fees) arising out of the Member's actions on behalf of the Company, except to the extent arising from the Member's gross negligence, willful misconduct, or fraud.")

    add_section(doc, 11, "MISCELLANEOUS")
    add_body(doc, f"11.1  Governing Law. This Agreement shall be governed by the laws of the State of {STATE}.")
    add_body(doc, "11.2  Entire Agreement. This Agreement constitutes the entire agreement of the Member with respect to the subject matter hereof and supersedes all prior agreements.")
    add_body(doc, "11.3  Amendment. This Agreement may be amended only by a written instrument signed by the Member.")
    add_body(doc, "11.4  Severability. If any provision of this Agreement is held invalid or unenforceable, the remaining provisions shall continue in full force and effect.")

    doc.add_paragraph()
    add_sig_block(doc, [MEMBER + " (Sole Member)"])

    doc.add_page_break()
    add_body(doc, "EXHIBIT A — INITIAL CAPITAL CONTRIBUTION")
    add_body(doc, f"Member: {MEMBER}")
    add_body(doc, "Initial Capital Contribution: $[AMOUNT] in cash / [DESCRIPTION OF NON-CASH CONTRIBUTION]")
    add_body(doc, "Membership Interest: 100%")

    doc.save("/home/user/Claude-/legal/01_LLC_Operating_Agreement.docx")
    print("✓ LLC Operating Agreement")


# ─── 2. IP Assignment Agreement ─────────────────────────────────────────────

def make_ip_assignment():
    doc = new_doc("IP Assignment Agreement")
    add_title(doc, "INTELLECTUAL PROPERTY ASSIGNMENT AGREEMENT")
    add_subtitle(doc, f"(Independent Contractor)")
    doc.add_paragraph()
    add_subtitle(doc, f"Effective Date: [DATE]")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"This Intellectual Property Assignment Agreement (\"Agreement\") is entered into as of the date set forth above, between {COMPANY}, a {STATE} limited liability company (\"Company\"), and the undersigned contractor (\"Contractor\").")

    add_section(doc, 1, "SERVICES")
    add_body(doc, "Contractor agrees to perform services for Company as described in a separate Statement of Work or engagement letter (\"Services\"). This Agreement applies to all Services performed, whether or not a formal Statement of Work is executed.")

    add_section(doc, 2, "ASSIGNMENT OF WORK PRODUCT")
    add_body(doc, "2.1  Assignment. Contractor hereby irrevocably assigns, transfers, and conveys to Company all right, title, and interest in and to all work product created, conceived, developed, or reduced to practice by Contractor in the performance of the Services (\"Work Product\"), including all intellectual property rights therein.")
    add_body(doc, "2.2  Work Product Defined. Work Product includes, without limitation: software code (source and object), scripts, algorithms, databases, schemas, designs, wireframes, graphics, documentation, inventions, discoveries, trade secrets, and any other materials or developments.")
    add_body(doc, "2.3  Works Made for Hire. To the maximum extent permitted by law, all Work Product is a \"work made for hire\" as defined under the U.S. Copyright Act. To the extent any Work Product is not deemed a work made for hire, Contractor's assignment in Section 2.1 shall apply.")
    add_body(doc, "2.4  Moral Rights. Contractor waives any moral rights or similar rights in the Work Product to the fullest extent permitted by law.")

    add_section(doc, 3, "PRE-EXISTING IP")
    add_body(doc, "3.1  Disclosure. Contractor shall disclose to Company in writing any pre-existing intellectual property (\"Pre-Existing IP\") that Contractor intends to incorporate into the Work Product before incorporating it.")
    add_body(doc, "3.2  License. Contractor grants Company a non-exclusive, perpetual, irrevocable, worldwide, royalty-free license to use, modify, and sublicense any Pre-Existing IP incorporated into the Work Product.")

    add_section(doc, 4, "CONFIDENTIALITY")
    add_body(doc, "Contractor agrees to keep all Company information, including product plans, source code, user data, business strategies, and financial information, strictly confidential during and after the engagement. Contractor shall not disclose Company confidential information to any third party without prior written consent.")

    add_section(doc, 5, "INDEPENDENT CONTRACTOR STATUS")
    add_body(doc, "Contractor is an independent contractor, not an employee, partner, or agent of Company. Contractor is responsible for all taxes on compensation received. Nothing in this Agreement creates an employment relationship.")

    add_section(doc, 6, "REPRESENTATIONS")
    add_body(doc, "Contractor represents and warrants that: (a) Contractor has full authority to enter into this Agreement; (b) the Work Product will be original and will not infringe any third-party intellectual property rights; (c) Contractor is not subject to any agreement that conflicts with this Agreement.")

    add_section(doc, 7, "GOVERNING LAW")
    add_body(doc, f"This Agreement shall be governed by the laws of the State of {STATE}. Any disputes shall be resolved in the courts of {STATE}.")

    add_section(doc, 8, "ENTIRE AGREEMENT")
    add_body(doc, "This Agreement constitutes the entire agreement between the parties regarding intellectual property and confidentiality and supersedes all prior understandings on these subjects.")

    add_sig_block(doc, [COMPANY, "Contractor"])
    doc.save("/home/user/Claude-/legal/02_IP_Assignment_Agreement.docx")
    print("✓ IP Assignment Agreement")


# ─── 3. NDA ─────────────────────────────────────────────────────────────────

def make_nda():
    doc = new_doc("NDA")
    add_title(doc, "MUTUAL NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph()
    add_subtitle(doc, f"Effective Date: [DATE]")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into between {COMPANY}, a {STATE} limited liability company (\"Company\"), and [COUNTERPARTY FULL LEGAL NAME], [an individual / a [STATE] entity] (\"Counterparty\"). Company and Counterparty are each referred to herein as a \"Party\" and collectively as the \"Parties.\"")

    add_body(doc, "The Parties wish to explore a potential business relationship (the \"Purpose\") and may disclose confidential information to each other in connection therewith. This Agreement governs the handling of such information.")

    add_section(doc, 1, "CONFIDENTIAL INFORMATION")
    add_body(doc, "\"Confidential Information\" means any non-public information disclosed by one Party (\"Disclosing Party\") to the other (\"Receiving Party\"), in any form, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure. This includes, without limitation: business plans, product roadmaps, source code, financial data, user data, technical specifications, and marketing strategies.")
    add_body(doc, "Confidential Information does not include information that: (a) is or becomes publicly available through no fault of the Receiving Party; (b) was rightfully known to the Receiving Party before disclosure; (c) is rightfully received from a third party without restriction; or (d) is independently developed by the Receiving Party without use of Confidential Information.")

    add_section(doc, 2, "OBLIGATIONS")
    add_body(doc, "Each Receiving Party agrees to: (a) hold the Disclosing Party's Confidential Information in strict confidence; (b) use it only for the Purpose; (c) not disclose it to any third party without prior written consent; (d) protect it with at least the same degree of care used for its own confidential information, but no less than reasonable care; (e) limit access to employees, contractors, and advisors who have a need to know and are bound by confidentiality obligations at least as protective as this Agreement.")

    add_section(doc, 3, "TERM")
    add_body(doc, "This Agreement is effective as of the date above and continues for two (2) years. Obligations with respect to trade secrets shall survive indefinitely.")

    add_section(doc, 4, "REQUIRED DISCLOSURE")
    add_body(doc, "If a Receiving Party is required by law or court order to disclose Confidential Information, it shall provide prompt written notice to the Disclosing Party (to the extent legally permitted) and cooperate with the Disclosing Party's efforts to seek a protective order.")

    add_section(doc, 5, "NO LICENSE")
    add_body(doc, "Nothing in this Agreement grants either Party any license, right, title, or interest in the other Party's Confidential Information except as necessary for the Purpose.")

    add_section(doc, 6, "RETURN OR DESTRUCTION")
    add_body(doc, "Upon request or termination of this Agreement, each Receiving Party shall promptly return or destroy all Confidential Information of the Disclosing Party and certify such destruction in writing.")

    add_section(doc, 7, "REMEDIES")
    add_body(doc, "Each Party acknowledges that breach of this Agreement may cause irreparable harm for which monetary damages would be inadequate, and that the non-breaching Party shall be entitled to seek equitable relief, including injunction, without bond.")

    add_section(doc, 8, "GOVERNING LAW")
    add_body(doc, f"This Agreement shall be governed by the laws of the State of {STATE}.")

    add_section(doc, 9, "ENTIRE AGREEMENT")
    add_body(doc, "This Agreement is the entire agreement between the Parties regarding confidentiality for the Purpose and supersedes all prior agreements on this subject.")

    add_sig_block(doc, [COMPANY, "Counterparty"])
    doc.save("/home/user/Claude-/legal/03_NDA.docx")
    print("✓ NDA")


# ─── 4. Plainly Terms of Service ────────────────────────────────────────────

def make_plainly_tos():
    doc = new_doc("Plainly ToS")
    add_title(doc, "PLAINLY — TERMS OF SERVICE")
    add_subtitle(doc, f"Operated by {COMPANY}")
    add_subtitle(doc, f"Last Updated: {TODAY}")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"These Terms of Service (\"Terms\") govern your access to and use of the Plainly mobile application and related services (\"Service\"), operated by {COMPANY} (\"we,\" \"us,\" or \"our\"). By creating an account or using the Service, you agree to these Terms. If you do not agree, do not use the Service.")

    add_section(doc, 1, "ELIGIBILITY")
    add_body(doc, "You must be at least 18 years of age to use the Service. By using the Service, you represent that you are 18 or older and have the legal capacity to enter into these Terms. The Service is not directed to persons under 13, and we do not knowingly collect data from children under 13.")

    add_section(doc, 2, "NOT FINANCIAL ADVICE")
    add_body(doc, "IMPORTANT: Plainly is an educational and informational tool only. Nothing in the Service — including AI coach responses, lesson content, personalized insights, or any other feature — constitutes financial advice, investment advice, tax advice, legal advice, or any regulated financial service.")
    add_body(doc, "We are not a registered investment advisor, broker-dealer, bank, or financial planner. The AI coach is an automated tool that provides general educational information based on inputs you provide. It does not take into account your full financial situation and is not a substitute for professional financial advice. Always consult a licensed financial advisor, accountant, or attorney before making financial decisions.")
    add_body(doc, "Any references to third-party financial products, institutions, or services are informational only and do not constitute endorsement.")

    add_section(doc, 3, "ACCOUNT REGISTRATION")
    add_body(doc, "3.1  You must provide accurate, complete information when creating an account. You are responsible for maintaining the confidentiality of your credentials and for all activity under your account.")
    add_body(doc, "3.2  You must notify us immediately at legal@plainly.app of any unauthorized use of your account.")
    add_body(doc, "3.3  We reserve the right to suspend or terminate accounts that violate these Terms.")

    add_section(doc, 4, "BANK ACCOUNT CONNECTION (PLAID)")
    add_body(doc, "4.1  The Service may allow you to connect your bank or financial accounts using Plaid Inc. (\"Plaid\"). By connecting accounts, you agree to Plaid's End User Privacy Policy and Terms of Service in addition to these Terms.")
    add_body(doc, "4.2  We receive read-only access to account information (balances, transactions) to personalize your experience. We do not initiate transactions, hold funds, or have write access to your accounts.")
    add_body(doc, "4.3  You may disconnect your bank account at any time through app settings. Disconnection does not automatically delete historical data already synced.")

    add_section(doc, 5, "SUBSCRIPTION AND PAYMENT")
    add_body(doc, "5.1  Plainly offers a free tier and optional paid subscription (\"Premium\"). Premium features are described within the app and are subject to change.")
    add_body(doc, "5.2  Subscriptions are billed on a recurring basis (monthly or annual) through the applicable app store (Apple App Store or Google Play). All billing is handled by the app store; we do not directly process payment card information.")
    add_body(doc, "5.3  Subscriptions automatically renew unless cancelled at least 24 hours before the renewal date. Manage or cancel subscriptions through your app store account settings.")
    add_body(doc, "5.4  We do not offer refunds except as required by applicable law or app store policy.")

    add_section(doc, 6, "ACCEPTABLE USE")
    add_body(doc, "You agree not to: (a) use the Service for any unlawful purpose; (b) attempt to reverse-engineer, decompile, or extract source code from the Service; (c) transmit malware, viruses, or harmful code; (d) scrape, harvest, or collect data from the Service; (e) impersonate any person or entity; (f) use the Service to provide financial services to third parties.")

    add_section(doc, 7, "INTELLECTUAL PROPERTY")
    add_body(doc, f"The Service and all content therein — including software, design, text, graphics, lesson content, and the Plainly brand — are owned by {COMPANY} or licensed to us. You receive a limited, non-exclusive, non-transferable license to use the Service for personal, non-commercial purposes. No other rights are granted.")

    add_section(doc, 8, "AI COACH — LIMITATIONS")
    add_body(doc, "The AI coach is powered by a third-party large language model. It may produce inaccurate, incomplete, or outdated information. Do not rely on coach responses for consequential financial decisions. We do not warrant the accuracy, completeness, or fitness for any purpose of any AI-generated content.")

    add_section(doc, 9, "DISCLAIMER OF WARRANTIES")
    add_body(doc, "THE SERVICE IS PROVIDED \"AS IS\" AND \"AS AVAILABLE\" WITHOUT WARRANTIES OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT. WE DO NOT WARRANT THAT THE SERVICE WILL BE UNINTERRUPTED, ERROR-FREE, OR SECURE.")

    add_section(doc, 10, "LIMITATION OF LIABILITY")
    add_body(doc, f"TO THE MAXIMUM EXTENT PERMITTED BY LAW, {COMPANY.upper()} SHALL NOT BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, OR LOSS OF PROFITS OR DATA, ARISING OUT OF OR RELATED TO THESE TERMS OR YOUR USE OF THE SERVICE. OUR TOTAL LIABILITY SHALL NOT EXCEED THE GREATER OF: (A) THE AMOUNT YOU PAID FOR THE SERVICE IN THE 12 MONTHS PRECEDING THE CLAIM, OR (B) $100.")

    add_section(doc, 11, "GOVERNING LAW AND DISPUTE RESOLUTION")
    add_body(doc, f"11.1  These Terms are governed by the laws of the State of {STATE}, without regard to conflict of law principles.")
    add_body(doc, "11.2  Informal Resolution. Before filing a claim, you agree to contact us at legal@plainly.app to attempt informal resolution.")
    add_body(doc, "11.3  Arbitration. [ATTORNEY TO REVIEW: Consider binding arbitration clause and class action waiver appropriate for your jurisdiction and user base.]")

    add_section(doc, 12, "CHANGES TO TERMS")
    add_body(doc, "We may update these Terms from time to time. We will notify you of material changes via in-app notification or email. Continued use after the effective date constitutes acceptance of the updated Terms.")

    add_section(doc, 13, "TERMINATION")
    add_body(doc, "We may suspend or terminate your access at any time for violation of these Terms or for any business reason. You may delete your account at any time. Upon termination, your license to use the Service ends immediately.")

    add_section(doc, 14, "CONTACT")
    add_body(doc, f"For legal notices: {COMPANY}, {MEMBER_ADDRESS}. Email: {EMAIL_PLAINLY}")

    doc.save("/home/user/Claude-/legal/04_Plainly_Terms_of_Service.docx")
    print("✓ Plainly Terms of Service")


# ─── 5. Vinla Terms of Service ──────────────────────────────────────────

def make_vinla_tos():
    doc = new_doc("Vinla ToS")
    add_title(doc, "HEALTH-AI — TERMS OF SERVICE")
    add_subtitle(doc, f"Operated by {COMPANY}")
    add_subtitle(doc, f"Last Updated: {TODAY}")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"These Terms of Service (\"Terms\") govern your use of the Vinla mobile application and related services (\"Service\"), operated by {COMPANY} (\"we,\" \"us,\" or \"our\"). By using the Service you agree to these Terms.")

    add_section(doc, 1, "ELIGIBILITY")
    add_body(doc, "You must be at least 18 years of age to use the Service. By using the Service, you represent that you are 18 or older. The Service is not directed to children under 13.")

    add_section(doc, 2, "NOT MEDICAL ADVICE — CRITICAL DISCLAIMER")
    add_body(doc, "IMPORTANT: Vinla is a personal wellness tracking and informational tool only. Nothing in the Service — including AI-generated insights, nutritional analysis, exercise recommendations, sleep analysis, or any other feature — constitutes medical advice, clinical diagnosis, medical treatment, or healthcare services of any kind.")
    add_body(doc, "We are not a healthcare provider, medical device manufacturer, or licensed health professional. The AI-generated insights are automated outputs based on data you log and are not reviewed by medical professionals. They are not a substitute for advice from a qualified physician, dietitian, or other licensed healthcare provider.")
    add_body(doc, "NEVER disregard or delay seeking professional medical advice because of something you read in the Service. If you have or suspect a medical condition, consult a licensed healthcare provider immediately. In an emergency, call 911 or your local emergency services.")
    add_body(doc, "The Service is not a medical device and has not been evaluated or approved by the U.S. Food and Drug Administration (FDA) or any other regulatory authority.")

    add_section(doc, 3, "HEALTH DATA YOU PROVIDE")
    add_body(doc, "3.1  You voluntarily provide health-related data (food intake, water consumption, exercise activity, sleep duration and quality, weight, and similar wellness metrics) to the Service.")
    add_body(doc, "3.2  You represent that data you enter is your own personal wellness data and that you have the right to submit it.")
    add_body(doc, "3.3  You are solely responsible for the accuracy of the data you enter. Inaccurate data will produce inaccurate insights.")

    add_section(doc, 4, "AI INSIGHTS — LIMITATIONS")
    add_body(doc, "AI-generated insights are based solely on data you provide and general wellness information in the model's training. Insights: (a) may be inaccurate or incomplete; (b) do not account for your full health history or medical conditions; (c) are not personalized medical recommendations; (d) should not be acted upon without consulting a healthcare professional for medical matters.")

    add_section(doc, 5, "ACCOUNT REGISTRATION")
    add_body(doc, "You must provide accurate registration information and maintain the confidentiality of your credentials. You are responsible for all activity under your account. Notify us at legal@<your vinla domain> of any unauthorized access.")

    add_section(doc, 6, "SUBSCRIPTION AND PAYMENT")
    add_body(doc, "6.1  Vinla offers a free tier and optional paid subscription for premium features. Premium features are described in-app.")
    add_body(doc, "6.2  Subscriptions are billed through the Apple App Store or Google Play. We do not directly process payment card information.")
    add_body(doc, "6.3  Subscriptions auto-renew unless cancelled at least 24 hours before renewal. Manage cancellations through your app store account.")

    add_section(doc, 7, "ACCEPTABLE USE")
    add_body(doc, "You agree not to: (a) use the Service for any unlawful purpose; (b) submit false or misleading health data; (c) attempt to reverse-engineer the Service; (d) use the Service to diagnose, treat, or provide health services to third parties; (e) scrape or harvest data.")

    add_section(doc, 8, "INTELLECTUAL PROPERTY")
    add_body(doc, f"The Service and all content — including software, design, AI model configurations, and the Vinla brand — are owned by {COMPANY}. You receive a limited personal license to use the Service. No other rights are granted.")

    add_section(doc, 9, "THIRD-PARTY INTEGRATIONS")
    add_body(doc, "The Service may integrate with third-party platforms (e.g., Apple HealthKit, Google Fit, wearable devices). Such integrations are subject to the third party's terms and privacy policies. We are not responsible for third-party services.")

    add_section(doc, 10, "DISCLAIMER OF WARRANTIES")
    add_body(doc, "THE SERVICE IS PROVIDED \"AS IS\" WITHOUT WARRANTIES OF ANY KIND. WE EXPRESSLY DISCLAIM ALL WARRANTIES, INCLUDING WARRANTIES RELATED TO ACCURACY OF HEALTH DATA ANALYSIS, FITNESS FOR A PARTICULAR PURPOSE, AND MEDICAL APPROPRIATENESS.")

    add_section(doc, 11, "LIMITATION OF LIABILITY")
    add_body(doc, f"TO THE MAXIMUM EXTENT PERMITTED BY LAW, {COMPANY.upper()} SHALL NOT BE LIABLE FOR ANY DAMAGES ARISING FROM YOUR USE OF OR RELIANCE ON THE SERVICE, INCLUDING DAMAGES ARISING FROM HEALTH DECISIONS MADE IN RELIANCE ON AI-GENERATED INSIGHTS. OUR TOTAL LIABILITY SHALL NOT EXCEED THE GREATER OF: (A) AMOUNTS PAID BY YOU IN THE PRIOR 12 MONTHS, OR (B) $100.")

    add_section(doc, 12, "GOVERNING LAW AND DISPUTES")
    add_body(doc, f"These Terms are governed by the laws of the State of {STATE}. [ATTORNEY TO REVIEW: arbitration clause, class action waiver, venue.]")

    add_section(doc, 13, "CHANGES AND TERMINATION")
    add_body(doc, "We may update these Terms at any time with notice via the app or email. We may suspend or terminate your account for violation of these Terms.")

    add_section(doc, 14, "CONTACT")
    add_body(doc, f"Legal notices: {COMPANY}, {MEMBER_ADDRESS}. Email: {EMAIL_VINLA}")

    doc.save("/home/user/Claude-/legal/05_Vinla_Terms_of_Service.docx")
    print("✓ Vinla Terms of Service")


# ─── 6. Plainly Privacy Policy ──────────────────────────────────────────────

def make_plainly_privacy():
    doc = new_doc("Plainly Privacy Policy")
    add_title(doc, "PLAINLY — PRIVACY POLICY")
    add_subtitle(doc, f"Operated by {COMPANY}")
    add_subtitle(doc, f"Last Updated: {TODAY}")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"This Privacy Policy explains how {COMPANY} (\"we,\" \"us,\" \"our\") collects, uses, shares, and protects information in connection with the Plainly mobile application (\"Service\"). By using the Service, you agree to this Privacy Policy.")

    add_section(doc, 1, "INFORMATION WE COLLECT")
    add_subsection(doc, "1.1  Information You Provide Directly")
    add_body(doc, "• Account registration data (name, email address, password)")
    add_body(doc, "• Financial profile information you enter (income, debt, savings goals)")
    add_body(doc, "• AI coach conversation messages")
    add_body(doc, "• Lesson progress and responses")

    add_subsection(doc, "1.2  Financial Account Data (via Plaid)")
    add_body(doc, "If you connect a bank or financial account, we receive through Plaid: account balances, transaction history (read-only), account names and types. We do not receive your bank login credentials. We have read-only access and cannot initiate transactions.")

    add_subsection(doc, "1.3  Automatically Collected Data")
    add_body(doc, "• Device information (device type, OS version, app version)")
    add_body(doc, "• Usage data (features used, session duration, screen views)")
    add_body(doc, "• IP address and general location (city/country level)")
    add_body(doc, "• Crash reports and performance data")

    add_subsection(doc, "1.4  AI Coach Conversations")
    add_body(doc, "Messages you send to the AI coach are transmitted to Anthropic, Inc. (our AI service provider) for processing. Your conversation context — including financial profile data injected into the session — is included in each API call. Anthropic does not use API customer data to train its AI models under its standard API agreement. Conversations are retained by us for up to 12 months to provide session context; see Section 3.")

    add_section(doc, 2, "HOW WE USE YOUR INFORMATION")
    add_body(doc, "We use your information to:")
    add_body(doc, "• Provide, operate, and personalize the Service")
    add_body(doc, "• Generate AI coach responses and personalized insights")
    add_body(doc, "• Track your lesson progress and send relevant nudges")
    add_body(doc, "• Process subscription payments (via app store)")
    add_body(doc, "• Send transactional emails and in-app notifications")
    add_body(doc, "• Improve the Service through analytics")
    add_body(doc, "• Detect fraud, abuse, and security threats")
    add_body(doc, "• Comply with legal obligations")
    add_body(doc, "We do not sell your personal information. We do not use your financial data for advertising purposes.")

    add_section(doc, 3, "SHARING YOUR INFORMATION")
    add_body(doc, "We share your information only with:")
    add_body(doc, "• Plaid Inc. — to connect financial accounts (subject to Plaid's privacy policy)")
    add_body(doc, "• Anthropic, Inc. (service provider) — to process AI coach conversations on our behalf. Anthropic acts as a data processor, not a data controller. Your data is not sold to Anthropic and is not used to train Anthropic's AI models. Anthropic's privacy policy: anthropic.com/privacy")
    add_body(doc, "• Cloud infrastructure providers (e.g., AWS, Google Cloud) — to host the Service")
    add_body(doc, "• Analytics providers — for aggregated, non-identifiable usage analytics")
    add_body(doc, "• Law enforcement or courts — when required by law or to protect rights and safety")
    add_body(doc, "• Successors — in connection with a merger, acquisition, or asset sale (with notice to you)")
    add_body(doc, "We require all third-party service providers to handle your data in accordance with applicable privacy law and our data processing agreements.")

    add_section(doc, 4, "DATA RETENTION")
    add_body(doc, "• Active account data: retained for the life of your account")
    add_body(doc, "• AI coach conversations: retained for [X] months to provide context")
    add_body(doc, "• Financial transaction data from Plaid: retained for [X] months")
    add_body(doc, "• Account data after deletion: purged within 30 days, except as required for legal compliance")

    add_section(doc, 5, "YOUR RIGHTS AND CHOICES")
    add_body(doc, "Depending on your jurisdiction, you may have the right to:")
    add_body(doc, "• Access a copy of your personal data")
    add_body(doc, "• Correct inaccurate data")
    add_body(doc, "• Request deletion of your account and data")
    add_body(doc, "• Disconnect your bank account at any time in app settings")
    add_body(doc, "• Opt out of marketing communications")
    add_body(doc, "• Data portability (export your data in a machine-readable format)")
    add_body(doc, f"To exercise these rights, contact us at {EMAIL_PLAINLY}. We will respond within 30 days.")

    add_section(doc, 6, "CALIFORNIA RESIDENTS (CCPA)")
    add_body(doc, "California residents have rights under the California Consumer Privacy Act (CCPA), including the right to know, delete, and opt out of sale of personal information. We do not sell personal information. To exercise CCPA rights, contact us at the address below. We will not discriminate against you for exercising your rights.")

    add_section(doc, 7, "SECURITY")
    add_body(doc, "We implement industry-standard security measures including encryption in transit (TLS), encryption at rest, access controls, and regular security reviews. No system is perfectly secure; we cannot guarantee absolute security.")

    add_section(doc, 8, "CHILDREN'S PRIVACY")
    add_body(doc, "The Service is not directed to persons under 18 (or 13 where applicable). We do not knowingly collect personal information from children under 13. If you believe a child has provided data, contact us immediately.")

    add_section(doc, 9, "CHANGES TO THIS POLICY")
    add_body(doc, "We may update this Privacy Policy. We will notify you of material changes via in-app notification or email at least 30 days before the effective date. Continued use constitutes acceptance.")

    add_section(doc, 10, "CONTACT")
    add_body(doc, f"{COMPANY} | {MEMBER_ADDRESS} | {EMAIL_PLAINLY}")

    doc.save("/home/user/Claude-/legal/06_Plainly_Privacy_Policy.docx")
    print("✓ Plainly Privacy Policy")


# ─── 7. Vinla Privacy Policy ────────────────────────────────────────────

def make_vinla_privacy():
    doc = new_doc("Vinla Privacy Policy")
    add_title(doc, "HEALTH-AI — PRIVACY POLICY")
    add_subtitle(doc, f"Operated by {COMPANY}")
    add_subtitle(doc, f"Last Updated: {TODAY}")
    doc.add_paragraph()
    add_attorney_note(doc)

    add_body(doc, f"This Privacy Policy explains how {COMPANY} (\"we,\" \"us,\" \"our\") collects, uses, shares, and protects information through the Vinla mobile application (\"Service\"). Health data is especially sensitive — we treat it with the highest level of care. By using the Service, you agree to this Privacy Policy.")

    add_section(doc, 1, "INFORMATION WE COLLECT")
    add_subsection(doc, "1.1  Health and Wellness Data You Log")
    add_body(doc, "• Food and nutrition data (meals, calories, macronutrients)")
    add_body(doc, "• Water intake")
    add_body(doc, "• Exercise and physical activity (type, duration, intensity)")
    add_body(doc, "• Sleep duration and quality ratings")
    add_body(doc, "• Body weight and body measurements (if you choose to log)")
    add_body(doc, "• Any other wellness metrics you voluntarily enter")

    add_subsection(doc, "1.2  Account Data")
    add_body(doc, "• Name, email address, and password")
    add_body(doc, "• Age, sex, height, and weight (for baseline calculations — optional)")
    add_body(doc, "• Health goals you set (e.g., target calories, sleep hours)")

    add_subsection(doc, "1.3  AI Interaction Data")
    add_body(doc, "Messages and queries you submit to the AI insight engine are transmitted to Anthropic, Inc. (our AI service provider) for processing. Your health data context included in the session is part of this transmission. Anthropic does not use API customer data to train its AI models. See Section 3 for full details.")

    add_subsection(doc, "1.4  Third-Party Health Platform Data (if connected)")
    add_body(doc, "If you connect Apple HealthKit, Google Fit, or a wearable device, we may receive data from those platforms with your explicit permission. Each integration is governed by that platform's own privacy policy.")

    add_subsection(doc, "1.5  Automatically Collected Data")
    add_body(doc, "• Device and OS information")
    add_body(doc, "• App usage patterns and session data")
    add_body(doc, "• Crash reports")

    add_section(doc, 2, "HOW WE USE YOUR INFORMATION")
    add_body(doc, "• Provide wellness tracking, trend analysis, and AI-generated insights")
    add_body(doc, "• Personalize your experience based on logged data")
    add_body(doc, "• Generate in-app notifications and insights")
    add_body(doc, "• Improve Service features through aggregated analytics")
    add_body(doc, "• Process premium subscription transactions")
    add_body(doc, "• Ensure security and prevent fraud")
    add_body(doc, "• Comply with legal obligations")
    add_body(doc, "We do NOT: sell your health data; use your health data for advertising; share identifiable health data with insurance companies, employers, or healthcare providers; use your data to train AI models without explicit consent.")

    add_section(doc, 3, "SHARING YOUR INFORMATION")
    add_body(doc, "• Anthropic, Inc. (service provider) — AI insight generation. Your health logging context and conversation data is transmitted to Anthropic to generate responses. Anthropic acts as a data processor only; your health data is not used to train AI models and is not sold. See anthropic.com/privacy for Anthropic's data handling policy.")
    add_body(doc, "• Cloud infrastructure providers — to host and secure the Service")
    add_body(doc, "• Analytics providers — aggregated, de-identified data only")
    add_body(doc, "• Apple / Google — subscription billing only")
    add_body(doc, "• Law enforcement — only when legally required or to prevent imminent harm")
    add_body(doc, "• Business successors — in a merger/acquisition (with advance notice)")
    add_body(doc, "We will never sell, rent, or share your identifiable health data with third parties for their own purposes.")

    add_section(doc, 4, "HEALTH DATA — SPECIAL PROTECTIONS")
    add_body(doc, "4.1  HIPAA Status. The Service is a general wellness application and is not a HIPAA-covered entity. We are not a healthcare provider or health plan. HIPAA does not apply to this Service. However, we voluntarily apply heightened protections to your health data.")
    add_body(doc, "4.2  Encryption. All health data is encrypted in transit (TLS 1.2+) and at rest (AES-256 or equivalent).")
    add_body(doc, "4.3  Access Controls. Access to health data is restricted to personnel who require it for Service operations.")
    add_body(doc, "4.4  No Sale of Health Data. We will never sell your health or biometric data under any circumstances.")
    add_body(doc, "4.5  State Biometric Laws. [ATTORNEY TO REVIEW: assess applicability of Illinois BIPA, Texas CUBI, Washington My Health MY Data Act, and similar state laws based on your user base and data types collected.]")

    add_section(doc, 5, "DATA RETENTION")
    add_body(doc, "• Active account data: retained for the life of your account")
    add_body(doc, "• Health log data: retained for [X] years to enable trend analysis")
    add_body(doc, "• AI insight conversation data: retained for [X] months")
    add_body(doc, "• Data after account deletion: purged within 30 days, except as required by law")

    add_section(doc, 6, "YOUR RIGHTS AND CHOICES")
    add_body(doc, "You may at any time:")
    add_body(doc, "• Access and review all health data you have logged")
    add_body(doc, "• Correct or delete individual log entries")
    add_body(doc, "• Export your complete health data in a portable format")
    add_body(doc, "• Delete your account and all associated data")
    add_body(doc, "• Revoke access to connected health platforms (Apple Health, Google Fit)")
    add_body(doc, "• Opt out of non-essential notifications")
    add_body(doc, f"To exercise data rights: {EMAIL_VINLA}. We respond within 30 days.")

    add_section(doc, 7, "CALIFORNIA RESIDENTS (CCPA)")
    add_body(doc, "California residents have the right to know, delete, and opt out of sale of personal information. We do not sell personal information or health data. Contact us to exercise your rights. We will not discriminate against you.")

    add_section(doc, 8, "SECURITY")
    add_body(doc, "We implement encryption, access controls, regular security assessments, and secure development practices. Despite these measures, no system is 100% secure. Notify us immediately at legal@<your vinla domain> of any suspected breach.")

    add_section(doc, 9, "CHILDREN'S PRIVACY")
    add_body(doc, "The Service is not directed to persons under 18. We do not knowingly collect health data from children under 13. Contact us immediately if you believe a child's data has been submitted.")

    add_section(doc, 10, "CHANGES")
    add_body(doc, "We will provide 30 days advance notice of material changes via in-app notification or email.")

    add_section(doc, 11, "CONTACT")
    add_body(doc, f"{COMPANY} | {MEMBER_ADDRESS} | {EMAIL_VINLA}")

    doc.save("/home/user/Claude-/legal/07_Vinla_Privacy_Policy.docx")
    print("✓ Vinla Privacy Policy")


# ─── Run all ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating legal documents...")
    make_operating_agreement()
    make_ip_assignment()
    make_nda()
    make_plainly_tos()
    make_vinla_tos()
    make_plainly_privacy()
    make_vinla_privacy()
    print("\nAll 7 documents generated in /legal/")
