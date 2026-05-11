#!/usr/bin/env python3
"""
Generate the Plainly Digital InfoSec Policy as a polished DOCX from the
markdown source at legal/04_Information_Security_Policy.md.

Run with:
    /tmp/docxenv/bin/python legal/generate_infosec_policy.py

Output:
    legal/04_Information_Security_Policy.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "04_Information_Security_Policy.md"
OUT = ROOT / "04_Information_Security_Policy.docx"

BODY_FONT = "Calibri"
BODY_SIZE = 11
TITLE_SIZE = 22
SECTION_SIZE = 14
SUBSECTION_SIZE = 12
TABLE_HEADER_BG = "1F3A5F"   # dark blue
GRAY_TEXT = RGBColor(0x44, 0x44, 0x44)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    return doc


def add_runs(paragraph, text: str, *, bold: bool = False, color: RGBColor | None = None):
    """Render inline markdown for bold (**...**) into runs on a paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)
        if color is not None:
            run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(TITLE_SIZE)
    run.font.color.rgb = ACCENT


def add_metadata_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text, color=GRAY_TEXT)


def add_section(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(SECTION_SIZE)
    run.font.color.rgb = ACCENT


def add_subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(SUBSECTION_SIZE)


def add_body(doc, text):
    p = doc.add_paragraph()
    add_runs(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    # Header row
    for col_idx, label in enumerate(header):
        cell = table.rows[0].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, TABLE_HEADER_BG)

    # Body rows
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_runs(para, value)
    doc.add_paragraph()


def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    p_pr.append(pBdr)


# ─── Markdown parser (just enough for THIS document) ────────────────────────

def parse_markdown(md: str):
    """
    Yield (kind, payload) tuples:
        ("title", "...")
        ("section", "...")
        ("subsection", "...")
        ("body", "...")
        ("bullet", "...")
        ("table", (header_cells, [row_cells, ...]))
        ("hr", None)
        ("blank", None)
    """
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            yield ("blank", None)
            i += 1
            continue
        if line.startswith("# ") and i == 0:
            yield ("title", line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            yield ("section", line[3:].strip())
            i += 1
            continue
        if line.startswith("### "):
            yield ("subsection", line[4:].strip())
            i += 1
            continue
        if line.strip() == "---":
            yield ("hr", None)
            i += 1
            continue
        if line.startswith("- "):
            yield ("bullet", line[2:].strip())
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            # Markdown table: header, separator, then rows until non-pipe line
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            yield ("table", (header_cells, rows))
            i = j
            continue
        # Default: body paragraph (collapse consecutive non-empty body lines)
        body_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (
                not nxt
                or nxt.startswith("# ")
                or nxt.startswith("## ")
                or nxt.startswith("### ")
                or nxt.startswith("- ")
                or nxt.startswith("|")
                or nxt.strip() == "---"
            ):
                break
            body_lines.append(nxt)
            i += 1
        yield ("body", " ".join(body_lines).strip())


def render():
    md = SRC.read_text(encoding="utf-8")
    doc = new_doc()

    # Track first few lines under the title — render as right-aligned metadata
    # (Effective date / Owner / Review cadence / Governing law) instead of bullets.
    in_metadata = False
    metadata_consumed = 0

    for kind, payload in parse_markdown(md):
        if kind == "title":
            add_title(doc, payload)
            # The metadata that follows is body lines starting with **
            in_metadata = True
            continue

        if in_metadata and kind == "body" and payload.startswith("**"):
            add_metadata_line(doc, payload)
            metadata_consumed += 1
            continue
        if in_metadata and (kind in ("hr", "section") or metadata_consumed >= 6):
            in_metadata = False
            # let normal handling take over

        if kind == "section":
            add_section(doc, payload)
        elif kind == "subsection":
            add_subsection(doc, payload)
        elif kind == "body":
            add_body(doc, payload)
        elif kind == "bullet":
            add_bullet(doc, payload)
        elif kind == "table":
            header, rows = payload
            add_table(doc, header, rows)
        elif kind == "hr":
            add_hr(doc)
        elif kind == "blank":
            # Trailing blank lines from the parser — let the body's
            # space_after handle the gap, no need to add explicit blanks.
            pass

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    render()
