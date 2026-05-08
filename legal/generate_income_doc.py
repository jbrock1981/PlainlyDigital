"""
Generate income projection document — possible revenue at each phase.
Covers Health-AI, Plainly, portfolio effects, and supplemental income streams.
Three scenarios: Conservative, Moderate, Optimistic.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TODAY = "March 2026"
DARK   = "1a1a2e"
TEAL   = "00b4d8"
ORANGE = "e76f51"
GREEN  = "2d6a4f"
PURPLE = "6d3a9c"
GOLD   = "c9a227"
GRAY   = "666666"


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin   = Inches(1.15)
        s.right_margin  = Inches(1.15)
    doc.styles['Normal'].font.name = 'Georgia'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(20)
    r.font.name = 'Georgia'; r.font.color.rgb = rgb(DARK)


def add_center(doc, text, size=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = 'Georgia'
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = rgb(color)
    return p


def add_rule(doc, color_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12')
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), color_hex)
    pBdr.append(b); pPr.append(pBdr)


def add_phase_header(doc, label, title, subtitle, color_hex):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}  —  ")
    r1.bold = True; r1.font.size = Pt(9)
    r1.font.name = 'Georgia'; r1.font.color.rgb = rgb(color_hex)
    r2 = p.add_run(title.upper())
    r2.bold = True; r2.font.size = Pt(13)
    r2.font.name = 'Georgia'; r2.font.color.rgb = rgb(DARK)
    p2 = doc.add_paragraph(subtitle)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    for r in p2.runs:
        r.font.name = 'Georgia'; r.font.size = Pt(10); r.italic = True
        r.font.color.rgb = rgb(GRAY)
    add_rule(doc, color_hex)


def add_body(doc, text, indent=0.0, size=10.5, italic=False, color=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(4)
    for r in p.runs:
        r.font.name = 'Georgia'; r.font.size = Pt(size)
        r.italic = italic
        if color: r.font.color.rgb = rgb(color)
    return p


def add_note(doc, text):
    add_body(doc, f"↳  {text}", indent=0.25, size=9.5, italic=True, color="555555")


def shd_cell(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, col_widths=None, header_color=DARK):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shd_cell(hdr[i], header_color)
        for para in hdr[i].paragraphs:
            for r in para.runs:
                r.bold = True; r.font.name = 'Georgia'
                r.font.size = Pt(10); r.font.color.rgb = rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for r in para.runs:
                    r.font.name = 'Georgia'; r.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                shd_cell(cells[ci], "f5f5f5")
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return t


def add_total_row(t, cells_data, color_hex=GREEN):
    row = t.add_row()
    for i, val in enumerate(cells_data):
        row.cells[i].text = str(val)
        shd_cell(row.cells[i], color_hex)
        for para in row.cells[i].paragraphs:
            for r in para.runs:
                r.bold = True; r.font.name = 'Georgia'
                r.font.size = Pt(10); r.font.color.rgb = rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER


def add_bullet(doc, title, detail, indent=0.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"• {title}: ")
    r1.bold = True; r1.font.name = 'Georgia'; r1.font.size = Pt(10.5)
    r2 = p.add_run(detail)
    r2.font.name = 'Georgia'; r2.font.size = Pt(10.5)


def scenario_box(doc, label, conservative, moderate, optimistic, color_hex):
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    labels = [label, "Conservative", "Moderate", "Optimistic"]
    for i, cell in enumerate(t.rows[0].cells):
        cell.text = labels[i]
        shd_cell(cell, color_hex)
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True; r.font.name = 'Georgia'
                r.font.size = Pt(10); r.font.color.rgb = rgb("ffffff")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vals = ["Monthly Revenue (net of App Store fee)", conservative, moderate, optimistic]
    for i, cell in enumerate(t.rows[1].cells):
        cell.text = vals[i]
        for para in cell.paragraphs:
            for r in para.runs:
                r.font.name = 'Georgia'; r.font.size = Pt(10.5)
                r.bold = (i > 0)
                if i > 0: r.font.color.rgb = rgb(GREEN)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def make_income_doc():
    doc = new_doc()

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_title(doc, "POSSIBLE INCOME BY PHASE")
    add_center(doc, "Revenue Projections — LLC Launch through Multi-App Portfolio", size=12, italic=True, color=GRAY)
    add_center(doc, "Health-AI  →  Plainly  →  Future Apps", size=11, bold=True, color=DARK)
    add_center(doc, f"Prepared {TODAY}", size=9, italic=True, color="999999")
    doc.add_paragraph()
    add_center(doc,
        "Three scenarios are shown throughout: Conservative (organic growth, no marketing spend), "
        "Moderate (some paid acquisition or community building), and Optimistic (strong PMF, "
        "word of mouth or press coverage). All revenue figures are net of App Store/Play Store "
        "15% commission. These are projections, not guarantees.",
        size=10, italic=True, color="555555")
    doc.add_paragraph()

    # ── KEY ASSUMPTIONS BOX ────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("KEY ASSUMPTIONS")
    r.bold = True; r.font.name = 'Georgia'; r.font.size = Pt(11)
    r.font.color.rgb = rgb(DARK)
    add_rule(doc, DARK)

    assumptions = [
        ("App Store commission", "15% on first $1M/year revenue (Apple Small Business + Google reduced fee program)"),
        ("Subscription pricing", "Health-AI: $9.99/mo or $79.99/yr  |  Plainly: $7.99/mo or $59.99/yr"),
        ("Freemium conversion rate", "Industry average: 2–5% of free users convert to paid"),
        ("Churn rate", "Monthly subscription churn: 5–8% is typical for consumer apps"),
        ("Growth method", "Organic/content marketing assumed for Conservative; some paid ads for Moderate/Optimistic"),
        ("'Net revenue' throughout", "After 15% App Store commission, before taxes and expenses"),
    ]
    for title, detail in assumptions:
        add_bullet(doc, title, detail)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 0 — LLC SETUP
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "PHASE 0", "LLC Formation", "Weeks 1–4. No apps live yet.", DARK)

    add_body(doc, "No direct app revenue during this phase. However, this window is valuable for "
             "generating pre-launch income and building the financial runway you'll need.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Income Opportunities During Phase 0:")
    r.bold = True; r.font.name = 'Georgia'; r.font.size = Pt(11)

    add_bullet(doc, "Freelance / consulting work",
        "Use your existing skills (development, design, marketing) to generate income while building. "
        "Even 10 hours/week at $50–$150/hr = $2,000–$6,000/month to fund the business.")
    add_bullet(doc, "Pre-launch email waitlist",
        "Build a landing page for Health-AI and collect emails. A list of 500–1,000 signups "
        "before launch dramatically improves Day 1 App Store ranking and validates the idea.")
    add_bullet(doc, "Content creation (YouTube, TikTok, X/Twitter)",
        "Document the build publicly — 'building in public' consistently drives early users. "
        "No immediate income, but builds an audience that converts at launch.")
    add_bullet(doc, "Grants & accelerator programs",
        "NSF SBIR Phase I grants: up to $275,000 for innovative tech (health and fintech qualify). "
        "Y Combinator, Techstars, and fintech-specific accelerators accept pre-product companies.")

    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    for i, txt in enumerate(["PHASE 0 INCOME", "Conservative", "Moderate", "Optimistic"]):
        t0.rows[0].cells[i].text = txt
        shd_cell(t0.rows[0].cells[i], DARK)
        for para in t0.rows[0].cells[i].paragraphs:
            for r in para.runs:
                r.bold=True; r.font.name='Georgia'; r.font.size=Pt(10); r.font.color.rgb=rgb("ffffff")
            para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for i, txt in enumerate(["Monthly (freelance/consulting)", "$0–$2,000/mo", "$2,000–$5,000/mo", "$5,000–$10,000/mo"]):
        t0.rows[1].cells[i].text = txt
        for para in t0.rows[1].cells[i].paragraphs:
            for r in para.runs:
                r.font.name='Georgia'; r.font.size=Pt(10.5)
                if i > 0: r.bold=True; r.font.color.rgb=rgb(GREEN)
            para.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 1 — HEALTH-AI
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "PHASE 1", "Health-AI — Revenue Ramp",
        "Months 1–12 post-launch. Building your first user base.", TEAL)

    add_body(doc, "Health-AI monetizes through a freemium subscription model. "
             "Free users get basic logging and limited AI insights. Premium unlocks "
             "unlimited AI coaching, advanced trend analysis, and goal tracking. "
             "The AI differentiation is the primary conversion driver.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Pricing Structure:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    add_bullet(doc, "Free tier", "Basic health logging, 3 AI insights/month, standard goal templates")
    add_bullet(doc, "Premium Monthly", "$9.99/month — unlimited AI coaching, trend analysis, custom goals, export")
    add_bullet(doc, "Premium Annual", "$79.99/year (~33% savings) — best for retention and LTV")
    add_bullet(doc, "Net per subscriber (after 15% App Store fee)",
        "$8.49/mo on monthly plan  |  $67.99/yr on annual plan (~$5.67/mo equivalent)")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Projected User Growth & Revenue (Health-AI):")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    headers = ["Timeframe", "Total Downloads", "Free Users", "Paid Subscribers", "Net Monthly Revenue"]
    rows = [
        ["Month 1 (launch)",   "200–500",       "195–490",    "5–10",      "$42–$85"],
        ["Month 3",            "500–2,000",      "480–1,960",  "20–40",     "$170–$340"],
        ["Month 6",            "1,500–5,000",    "1,425–4,750","75–250",    "$637–$2,123"],
        ["Month 9",            "3,000–10,000",   "2,850–9,500","150–500",   "$1,274–$4,245"],
        ["Month 12 (Year 1)",  "5,000–20,000",   "4,750–19,000","250–1,000","$2,123–$8,490"],
    ]
    add_table(doc, headers, rows, col_widths=[1.3, 1.2, 1.2, 1.2, 1.3], header_color=TEAL)
    add_note(doc, "Assumes 5% freemium conversion rate, 70/30 split monthly/annual, 6% monthly churn. "
             "Conservative = left value, Optimistic = right value.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Additional Revenue Streams — Health-AI:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    add_bullet(doc, "B2B / Employer Wellness",
        "Sell Health-AI access to small employers as a wellness benefit. "
        "$5–$15/employee/month. Even 1 company with 20 employees = $100–$300/mo additional. "
        "Requires minimal extra development — just a team dashboard.")
    add_bullet(doc, "Health Coach Referral Partnerships",
        "Partner with certified health coaches who recommend the app to their clients. "
        "Revenue share: 20–30% of subscription for referred users. Low effort, compounding.")
    add_bullet(doc, "Premium one-time content packs",
        "Sell structured programs (30-day challenges, nutrition plans) as in-app purchases. "
        "$4.99–$9.99 one-time. Supplements subscription revenue.")
    add_bullet(doc, "App Store feature potential",
        "Health & Fitness is one of Apple's highest-featured categories. A single feature "
        "in 'New Apps We Love' or 'App of the Day' can drive 10,000–100,000 downloads in a week.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Year 1 Health-AI Revenue Summary:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    headers2 = ["Scenario", "Paid Subscribers (Mo 12)", "Avg MRR (Months 6–12)", "Year 1 Total Net Revenue"]
    rows2 = [
        ["Conservative (organic only)",      "250",   "$1,200/mo",  "$5,000–$9,000"],
        ["Moderate (some community/content)", "500",   "$2,800/mo",  "$14,000–$20,000"],
        ["Optimistic (press/App Store feature)","1,000","$5,500/mo", "$28,000–$45,000"],
    ]
    t2 = add_table(doc, headers2, rows2, col_widths=[2.2, 1.5, 1.5, 1.6], header_color=TEAL)
    add_note(doc, "Year 1 total is lower than MRR × 12 because revenue ramps — first 3 months are near-zero.")

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 2 — PLAINLY
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "PHASE 2", "Plainly — Financial Literacy Revenue",
        "Months 6–18 (launching ~6 months after Health-AI). Second income stream.", ORANGE)

    add_body(doc, "Plainly has a larger addressable market than Health-AI — financial stress affects nearly "
             "every adult. The app targets 18–35 year olds learning to budget, save, and invest. "
             "Plainly's B2B opportunity (credit unions, HR departments, schools) is particularly strong "
             "and can generate revenue faster than consumer subscriptions.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Pricing Structure:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    add_bullet(doc, "Free tier", "Core financial literacy lessons, basic budget tracker, 3 AI coach questions/month")
    add_bullet(doc, "Premium Monthly", "$7.99/month — unlimited AI coach, Plaid bank connection, full lesson library, goal tracking")
    add_bullet(doc, "Premium Annual", "$59.99/year (~37% savings)")
    add_bullet(doc, "Net per subscriber (after 15% fee)", "$6.79/mo monthly  |  $50.99/yr annual (~$4.25/mo equivalent)")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("B2B Revenue — The Bigger Opportunity:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    add_body(doc,
        "Financial literacy has massive B2B demand. Employers, credit unions, community banks, "
        "high schools, and colleges all need solutions. B2B contracts are larger, stickier, "
        "and don't depend on App Store discovery.", indent=0.2)

    add_bullet(doc, "Credit Unions & Community Banks",
        "Many are legally required to provide financial education. Site license: $500–$2,500/month. "
        "Even 1 credit union client = $500–$2,500 MRR instantly.")
    add_bullet(doc, "Employer Financial Wellness Programs",
        "$3–$10/employee/month. A 200-person company = $600–$2,000/month. "
        "HR departments actively budget for this — it reduces financial stress = higher productivity.")
    add_bullet(doc, "High Schools & Community Colleges",
        "Per-student licensing: $5–$15/student/semester. A single district = $5,000–$50,000/yr.")
    add_bullet(doc, "Nonprofit Financial Counseling Orgs",
        "Discounted licensing at $1,000–$3,000/year per organization. "
        "Lower revenue but strong credibility and referral network.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Projected Revenue — Plainly (Year 1 of Plainly's Launch):")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    headers3 = ["Stream", "Conservative", "Moderate", "Optimistic"]
    rows3 = [
        ["Consumer subscriptions (Mo 12)",         "$800/mo",    "$2,500/mo",  "$7,000/mo"],
        ["B2B — Credit unions / banks",            "$0",         "$1,500/mo",  "$5,000/mo"],
        ["B2B — Employer wellness",                "$0",         "$600/mo",    "$3,000/mo"],
        ["B2B — Schools / nonprofits",             "$0",         "$500/mo",    "$2,000/mo"],
        ["In-app one-time purchases",              "$100/mo",    "$300/mo",    "$800/mo"],
    ]
    t3 = add_table(doc, headers3, rows3, col_widths=[2.5, 1.3, 1.3, 1.3], header_color=ORANGE)
    add_total_row(t3, ["TOTAL MRR (Month 12 of Plainly)", "$900/mo", "$5,400/mo", "$17,800/mo"], ORANGE)
    doc.add_paragraph()
    add_note(doc,
        "B2B revenue requires active outreach — attending credit union conferences, "
        "LinkedIn outreach to HR Directors, and school district partnership efforts. "
        "It doesn't happen automatically but moves faster than consumer app growth.")

    # ════════════════════════════════════════════════════════════════════════
    # COMBINED VIEW — BOTH APPS
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "COMBINED", "Both Apps Running — Portfolio Revenue",
        "Month 12–18: Health-AI mature, Plainly ramping. Revenue compounds.", PURPLE)

    add_body(doc,
        "Once both apps are live and have 6–12 months of growth, your monthly revenue "
        "compounds. Expenses don't scale 2x — the shared infrastructure means your "
        "margin improves significantly as you add apps and users.")

    headers4 = ["Revenue Source", "Conservative", "Moderate", "Optimistic"]
    rows4 = [
        ["Health-AI subscriptions",           "$2,000/mo",  "$5,500/mo",  "$12,000/mo"],
        ["Health-AI B2B (employer wellness)", "$0",         "$500/mo",    "$3,000/mo"],
        ["Plainly subscriptions",             "$800/mo",    "$2,500/mo",  "$7,000/mo"],
        ["Plainly B2B (CU/banks/employers)",  "$500/mo",    "$2,600/mo",  "$10,000/mo"],
        ["Consulting / freelance",            "$2,000/mo",  "$1,000/mo",  "$0 (full-time on apps)"],
        ["Grants / awards",                   "$0",         "$1,000/mo",  "$2,500/mo"],
    ]
    t4 = add_table(doc, headers4, rows4, col_widths=[2.5, 1.3, 1.3, 1.3], header_color=PURPLE)
    add_total_row(t4, ["COMBINED GROSS MRR", "~$5,300/mo", "~$13,100/mo", "~$34,500/mo"], PURPLE)
    doc.add_paragraph()

    add_note(doc, "Less monthly infrastructure costs (~$111–$417/mo). Effective margin at these revenue levels: 95%+.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Annual Revenue Equivalent:")
    r.bold = True; r.font.name='Georgia'; r.font.size=Pt(11)

    add_bullet(doc, "Conservative", "~$63,600/year gross — modest but covers costs and founder salary at lean lifestyle")
    add_bullet(doc, "Moderate",     "~$157,200/year gross — solid indie founder income; hire a part-time contractor")
    add_bullet(doc, "Optimistic",   "~$414,000/year gross — fundable business; hire a small team")
    add_note(doc, "These are pre-tax gross figures. After self-employment tax and expenses, "
             "plan for 60–70% of gross as effective take-home in Conservative/Moderate scenarios.")

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 3+ — FUTURE APPS
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "PHASE 3+", "Future Apps — Compounding Returns",
        "Each additional app is nearly pure margin on top of existing infrastructure.", GOLD)

    add_body(doc,
        "By Phase 3, you have a proven playbook: build → launch → grow → repeat. "
        "Your shared backend, auth, design system, App Store accounts, and brand recognition "
        "mean each new app launches faster and costs far less to operate. "
        "This is the compound effect of building a studio rather than a single app.")

    add_bullet(doc, "App #3 incremental infra cost", "~$15–$112/year (domain + policy only)")
    add_bullet(doc, "App #3 time to launch",         "50–70% faster than App #1 — patterns are established")
    add_bullet(doc, "App #3 year 1 revenue potential", "$10,000–$60,000 net, depending on category")
    add_bullet(doc, "Portfolio effect",
        "Users of one app are warm leads for others. Cross-promotion between Health-AI and Plainly "
        "is natural — financial health and physical health are deeply connected.")
    add_bullet(doc, "Studio valuation",
        "A portfolio of 3+ profitable apps with combined MRR of $10K+ is sellable. "
        "App studios typically sell at 2–4x annual revenue. At $150K ARR, that's $300K–$600K.")

    # ════════════════════════════════════════════════════════════════════════
    # NON-APP INCOME STREAMS
    # ════════════════════════════════════════════════════════════════════════
    add_phase_header(doc, "SUPPLEMENTAL", "Other Income Streams (Any Phase)",
        "Income sources that run in parallel with app development.", GRAY)

    streams = [
        ("NSF SBIR / STTR Grants",
         "$50,000–$275,000 one-time",
         "Health-AI (digital health innovation) and Plainly (financial inclusion technology) both qualify. "
         "Phase I awards are non-dilutive — you keep full equity. Apply at seedfund.nsf.gov."),
        ("Fintech Accelerators",
         "$20,000–$150,000 + mentorship",
         "Visible Hands, Greenwood Capital, CFI (Center for Financial Inclusion), and "
         "CFPB Tech Sprint programs specifically fund financial literacy and inclusion tech."),
        ("Health Tech Grants",
         "$10,000–$100,000",
         "Robert Wood Johnson Foundation, Blue Shield Foundation, and state public health departments "
         "fund digital health tools. Non-dilutive. Requires application and reporting."),
        ("Content / Creator Income",
         "$500–$10,000+/month",
         "Document the build on YouTube or TikTok. Tutorial content (React Native, Expo, Claude API integration) "
         "monetizes through AdSense, sponsorships, and digital products. Builds audience for app launches."),
        ("Courses / Digital Products",
         "$1,000–$20,000 per launch",
         "Sell a course on building apps with Expo + Claude API. "
         "Or a Plainly financial literacy curriculum packaged for schools. "
         "One-time creation, recurring passive revenue."),
        ("Angel / Pre-Seed Investment",
         "$50,000–$500,000 (dilutive)",
         "Once you have any traction (100+ paying users), you're fundable for a pre-seed round. "
         "Health-AI and Plainly both operate in hot investment categories. "
         "Trade equity for runway to hire and grow faster."),
    ]

    for title, amount, detail in streams:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"• {title}  ")
        r1.bold = True; r1.font.name='Georgia'; r1.font.size=Pt(10.5)
        r2 = p.add_run(f"({amount})")
        r2.font.name='Georgia'; r2.font.size=Pt(10.5); r2.bold=True
        r2.font.color.rgb = rgb(GREEN)
        p2 = doc.add_paragraph(f"    {detail}")
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.space_after = Pt(6)
        for r in p2.runs:
            r.font.name='Georgia'; r.font.size=Pt(10)
            r.font.color.rgb = rgb("444444")

    # ════════════════════════════════════════════════════════════════════════
    # MASTER TIMELINE TABLE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("MASTER INCOME TIMELINE — CONSERVATIVE SCENARIO")
    r.bold=True; r.font.name='Georgia'; r.font.size=Pt(12); r.font.color.rgb=rgb(DARK)
    add_rule(doc, DARK)

    headers5 = ["Period", "Status", "Monthly Revenue", "Cumulative Revenue", "Notes"]
    rows5 = [
        ["Month 0–1",   "LLC forming",                "$0–$2,000",     "$0–$2,000",     "Consulting income only"],
        ["Month 2–4",   "Building Health-AI",         "$0–$2,000",     "$0–$8,000",     "Consulting + possible grant apps"],
        ["Month 5",     "Health-AI launches",         "$42–$2,000",    "$42–$10,000",   "First App Store revenue"],
        ["Month 6",     "Health-AI growing",          "$637–$3,000",   "$1,000–$16,000","Starting to build list for Plainly"],
        ["Month 9",     "Building Plainly",           "$1,274–$4,000", "$5,000–$30,000","Health-AI still growing"],
        ["Month 12",    "Plainly launches",           "$2,123–$5,000", "$12,000–$50,000","Two revenue streams begin"],
        ["Month 15",    "Plainly ramping + B2B",      "$3,500–$8,000", "$25,000–$80,000","First B2B deals possible"],
        ["Month 18",    "Both apps mature",           "$5,300–$13,000","$50,000–$130,000","Consider Phase 3 app"],
        ["Month 24",    "Portfolio (3 apps)",         "$8,000–$20,000","$100,000–$250,000","Sustainable indie studio"],
    ]
    add_table(doc, headers5, rows5, col_widths=[1.0, 1.5, 1.2, 1.4, 1.7], header_color=DARK)

    add_note(doc,
        "Month 24 projections assume Phase 3 app launched around month 18. "
        "Conservative scenario assumes zero paid marketing, organic App Store discovery only.")

    # ── Footer ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Revenue projections based on industry benchmarks for consumer subscription apps (2024–2026). "
        "Freemium conversion rates from Lenny's Newsletter, RevenueCat State of Subscription Apps 2025, "
        "and App Store analytics benchmarks. B2B pricing based on comparable fintech and health SaaS tools. "
        "All figures are estimates — actual results depend heavily on product quality, marketing, and market fit."
    )
    for r in p.runs:
        r.font.name='Georgia'; r.font.size=Pt(9); r.italic=True
        r.font.color.rgb=rgb("999999")

    doc.save("/home/user/Claude-/docs/income_projections.docx")
    print("✓ Income projections generated")


if __name__ == "__main__":
    make_income_doc()
